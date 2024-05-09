# -*- coding:utf-8 -*-
"""
@Author: Smallsea
"""
import tensorflow as tf
import numpy as np
import codecs
import pickle
import os
from datetime import datetime
from bert_base.train.models import create_model, InputFeatures
from bert_base.bert import tokenization, modeling
from bert_base.train.train_helper import get_args_parser
import re
import unicodedata  # 处理ASCii码的包
import json
import uuid
from docx import Document
import shutil
import sys
import time

sys.argv = [sys.argv[0]]
args = get_args_parser()
absp = os.path.abspath('..')  # 单独测试时
absp = os.path.abspath('.')  # 系统测试时
model_dir = f'{absp}/data/output'
bert_dir = f'{absp}/data/chinese_L-12_H-768_A-12'
# model_dir = r'F:/mysite_v1/data/output'
# bert_dir = r'F:/mysite_v1/data/chinese_L-12_H-768_A-12'
# if os.name == 'nt':
#     bert_path = 'D:/NER/chinese_L-12_H-768_A-12'
#     root_path = r'D:/NER/BERT_BiLSTM_CRF_NER'
# else:
#     bert_path = '/home/cxh18/teamwork/NER/chinese_L-12_H-768_A-12/'
#     root_path = '/home/cxh18/teamwork/NER/BERT_BiLSTM_CRF_NER'

is_training = False
use_one_hot_embeddings = False
batch_size = 1
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'
gpu_config = tf.ConfigProto()
gpu_config.gpu_options.allow_growth = True
sess = tf.Session(config=gpu_config)
model = None

global graph
input_ids_p, input_mask_p, label_ids_p, segment_ids_p = None, None, None, None

# print('checkpoint path:{}'.format(os.path.join(model_dir, "checkpoint")))
if not os.path.exists(os.path.join(model_dir, "checkpoint")):
    raise Exception("failed to get checkpoint. going to return ")

# 加载label->id的词典
with codecs.open(os.path.join(model_dir, 'label2id.pkl'), 'rb') as rf:
    label2id = pickle.load(rf)
    id2label = {value: key for key, value in label2id.items()}

with codecs.open(os.path.join(model_dir, 'label_list.pkl'), 'rb') as rf:
    label_list = pickle.load(rf)
num_labels = len(label_list) + 1

graph = tf.get_default_graph()
with graph.as_default():
    # print("going to restore checkpoint")
    # sess.run(tf.global_variables_initializer())
    input_ids_p = tf.placeholder(tf.int32, [batch_size, args.max_seq_length], name="input_ids")
    input_mask_p = tf.placeholder(tf.int32, [batch_size, args.max_seq_length], name="input_mask")

    bert_config = modeling.BertConfig.from_json_file(os.path.join(bert_dir, 'bert_config.json'))
    (total_loss, logits, trans, pred_ids) = create_model(
        bert_config=bert_config, is_training=False, input_ids=input_ids_p, input_mask=input_mask_p, segment_ids=None,
        labels=None, num_labels=num_labels, use_one_hot_embeddings=False, dropout_rate=1.0)

    saver = tf.train.Saver()
    saver.restore(sess, tf.train.latest_checkpoint(model_dir))

tokenizer = tokenization.FullTokenizer(vocab_file=os.path.join(bert_dir, 'vocab.txt'), do_lower_case=args.do_lower_case)


def extract_one_sample(files):
    # create_directory()
    start = datetime.now()
    file_name = os.path.splitext(files)[0]
    file_property = os.path.splitext(files)[1]
    DocxSplitTxt(file_name)
    print('文本分割耗费时间: {} sec'.format((datetime.now() - start).total_seconds()))

    if not file_name:
        print("not this docx, ignore...")
        return
    print("smart extract information...")

    def convert(line):
        feature = convert_single_example(0, line, label_list, args.max_seq_length, tokenizer, 'p')
        input_ids = np.reshape([feature.input_ids], (batch_size, args.max_seq_length))
        input_mask = np.reshape([feature.input_mask], (batch_size, args.max_seq_length))
        segment_ids = np.reshape([feature.segment_ids], (batch_size, args.max_seq_length))
        label_ids = np.reshape([feature.label_ids], (batch_size, args.max_seq_length))
        return input_ids, input_mask, segment_ids, label_ids

    # 生化指标.肝功能检查
    PATTERN_21_1 = r"([\u4e00-\u9fa5]*)[、|,|，|。|\.|\\|\/|~]?(肝病|肝功能不全|肝)[、|,|，|。|\.|\\|\/|~]?([\u4e00-\u9fa5]*)"
    # 生化指标.肾功能检查
    PATTERN_22_1 = r"([\u4e00-\u9fa5]*)[、|,|，|。|\.|\\|\/|~]?(肾病|肾功能不全|肾)[、|,|，|。|\.|\\|\/|~]?([\u4e00-\u9fa5]*)"
    # 生化指标.肝功能检查
    compiled_rule_21_1 = re.compile(PATTERN_21_1)
    # 生化指标.肾功能检查
    compiled_rule_22_1 = re.compile(PATTERN_22_1)

    global graph
    with graph.as_default():
        # print(id2label)
        path_instructions = f"{absp}/data/instructions_split/" + file_name
        if not os.path.isdir(path_instructions):
            print(f"not exist this dir：{file_name}")
            return
        start_all = datetime.now()
        for files in os.walk(path_instructions):
            # print(files)
            field_dict = {}
            field_dict["年龄"] = []
            others_field = {}
            others_field["用药间隔周期"] = []
            others_field["给药频率"] = []
            others_field["使用疗程"] = []
            others_field["每次给药剂量"] = []
            others_field["每公斤每次极量"] = []
            others_field["每天剂量"] = []
            others_field["给药目的"] = []
            others_field["每公斤每日剂量"] = []
            others_field["每公斤每日极量"] = []
            others_field["次极量"] = []
            others_field["日极量"] = []
            others_field["每公斤总剂量"] = []
            others_field["每公斤每分钟滴速"] = []
            others_field["滴速"] = []
            others_field["持续时间"] = []
            others_field["每平方米每次剂量"] = []
            others_field["每平方米每次极量"] = []
            others_field["每平方米每日剂量"] = []
            others_field["每平方米每日极量"] = []
            others_field["每平方米总剂量"] = []
            others_field["总剂量"] = []
            others_field["药品剂型"] = []
            others_field["给药时机"] = []
            others_field["给药途径"] = []

            DRUGNAME = ''
            node = {}
            node["病人资料.性别"] = ""  # 可能需要改成仿照年龄
            node["生化指标.肝功能检查.谷丙转氨酶"] = ""
            node["生化指标.肾功能检查.内生肌酐清除率"] = ""
            node["病人资料.年龄"] = {}
            node["病人资料.是否怀孕"] = ""
            node["病人资料.是否哺乳"] = ""
            node["病人资料.过敏药品列表"] = ""
            node["药嘱.给药途径"] = []
            alter_sign = False
            first_sign = True
            drug_json = {}
            NODE = {}
            NODE['用法用量'] = {}
            NODE["给药途径"] = {}
            NODE["相互作用"] = {}
            NODE["配伍"] = {}
            NODE["患者条件"] = {}
            NODE["孕产"] = {}
            NODE["过敏"] = {}
            NODE["生化指标"] = {}
            NODE["重复用药"] = {}
            for dd in files[2]:
                if first_sign:
                    # print(files[0])
                    first_sign = False
                if dd == "首行药品名称.txt":
                    DRUGNAME = search_drugName(files[0])
                    drug_json[f"{DRUGNAME}"] = NODE
                elif dd == '用法用量.txt' or dd == '用法与用量.txt' or dd == '用量用法.txt' or dd == '用量与用法.txt' or dd == '免疫程序和剂量.txt':
                    alter_sign = True
                    with open(f"{files[0]}/{dd}", 'r', encoding='utf-8') as txt:
                        while True:
                            reader = txt.readline()
                            if reader == "":
                                break
                            sentence = reader.replace("\n", "")
                            sentence = sentence.replace(" ", "")

                            # 先序列截断预测年龄
                            sentence_list = []
                            max_seq_length = args.max_seq_length
                            sign_split = True
                            while sign_split:
                                if len(sentence) >= max_seq_length - 1:
                                    tokens = sentence[0:(max_seq_length - 2)]
                                    sentence_list.append(tokens)
                                    sentence = sentence[max_seq_length - 2:]
                                else:
                                    sentence_list.append(sentence)
                                    sign_split = False
                            sign_firstSplit = True
                            reader0 = ""
                            phase_dict = {}
                            age_field = []
                            # print(f"sentence_list = {sentence_list}")
                            start = datetime.now()
                            last_key = ""
                            for sentence in sentence_list:
                                sentence_token = tokenizer.tokenize(sentence)
                                input_ids, input_mask, segment_ids, label_ids = convert(sentence_token)

                                feed_dict = {input_ids_p: input_ids, input_mask_p: input_mask}
                                # run session get current feed_dict result
                                pred_ids_result = sess.run([pred_ids], feed_dict)
                                pred_label_result = convert_id_to_label(pred_ids_result, id2label)
                                # 输出预测标签
                                # print(pred_label_result[0])

                                # 病人资料.年龄
                                if sign_firstSplit:
                                    phase_dict0, reader0, age_field0 = search_age(sentence, sentence_token,
                                                                                  pred_label_result[0])
                                    # print(f"phase_dict0 = {phase_dict0}")
                                    phase_dict.update(phase_dict0)
                                    age_field.extend(age_field0)
                                    sign_firstSplit = False
                                    if phase_dict:
                                        last_key = list(phase_dict.keys())[-1]
                                    # print(f"phase_dict0_1 = {phase_dict}")
                                else:
                                    phase_dict0, reader1, age_field0 = search_age(sentence, sentence_token,
                                                                                  pred_label_result[0])
                                    if not last_key:
                                        reader0 += reader1
                                        if phase_dict0:
                                            last_key = list(phase_dict0.keys())[-1]
                                    else:
                                        phase_dict[last_key] += reader1
                                        last_key = list(phase_dict.keys())[-1]
                                    # 句子累加
                                    for key in phase_dict0:
                                        if key in phase_dict.keys():
                                            phase_dict[key] += phase_dict0[key]
                                        else:
                                            phase_dict.update({key: phase_dict0[key]})
                                    age_field.extend(age_field0)
                                    # print(f"phase_dict1 = {phase_dict0}")
                                    # print(f"phase_dict1_1 = {phase_dict}")
                            # print(f"phase_dict = {phase_dict}")
                            # 去重
                            age_field = list(set(age_field))
                            # print(f"phase_dict = {phase_dict}")
                            field_dict["年龄"] = age_field
                            # print(f"年龄 = {age_field}")
                            else_sign = False  # 是否需要有else

                            print('年龄识别耗费时间: {} sec'.format((datetime.now() - start).total_seconds()))
                            start = datetime.now()
                            for sentence_split in phase_dict:
                                else_sign = True
                                # print(sentence_split)
                                sentence_split_token = tokenizer.tokenize(phase_dict[sentence_split])
                                # print(sentence_token)
                                # print('your input is:{}'.format(sentence_token))
                                input_ids, input_mask, segment_ids, label_ids = convert(sentence_split_token)

                                feed_dict = {input_ids_p: input_ids,
                                             input_mask_p: input_mask}
                                # run session get current feed_dict result
                                pred_ids_result = sess.run([pred_ids], feed_dict)
                                pred_label_result = convert_id_to_label(pred_ids_result, id2label)

                                node["病人资料.年龄"][sentence_split] = {}
                                node["药嘱.用药间隔周期"] = ""
                                node["药嘱.给药频率"] = ""
                                node["药嘱.使用疗程"] = ""
                                node["药嘱.每次给药剂量"] = ""
                                node["药嘱.每公斤每次剂量"] = ""
                                node["药嘱.每公斤每次极量"] = ""
                                node["药嘱.每天剂量"] = ""
                                node["药嘱.给药目的"] = ""
                                node["药嘱.每公斤每日剂量"] = ""

                                node["药嘱.每公斤每日极量"] = ""
                                node["药嘱.次极量"] = ""
                                node["药嘱.日极量"] = ""
                                node["药嘱.每公斤总剂量"] = ""
                                node["药嘱.每公斤每分钟滴速"] = ""
                                node["药嘱.滴速"] = ""
                                node["药嘱.持续时间"] = ""
                                node["药嘱.每平方米每次剂量"] = ""
                                node["药嘱.每平方米每次极量"] = ""
                                node["药嘱.每平方米每日剂量"] = ""
                                node["药嘱.每平方米每日极量"] = ""
                                node["药嘱.每平方米总剂量"] = ""
                                node["药嘱.总剂量"] = ""
                                node["药嘱.药品剂型"] = ""
                                node["药嘱.给药时机"] = ""

                                node["药嘱.用药间隔周期"], node["药嘱.给药频率"], node["药嘱.使用疗程"], \
                                node["药嘱.每次给药剂量"], node["药嘱.每公斤每次剂量"], node["药嘱.每公斤每次极量"], \
                                node["药嘱.每天剂量"], node["药嘱.给药目的"], node["药嘱.每公斤每日剂量"], \
                                node["药嘱.每公斤每日极量"], node["药嘱.次极量"], node["药嘱.日极量"], \
                                node["药嘱.每公斤总剂量"], node["药嘱.每公斤每分钟滴速"], node["药嘱.滴速"], \
                                node["药嘱.持续时间"], node["药嘱.每平方米每次剂量"], node["药嘱.每平方米每次极量"], \
                                node["药嘱.每平方米每日剂量"], node["药嘱.每平方米每日极量"], node["药嘱.每平方米总剂量"], \
                                node["药嘱.总剂量"], node["药嘱.药品剂型"], node["药嘱.给药时机"], dosingWay = search_field(
                                    sentence_split_token, pred_label_result[0], others_field)
                                node["药嘱.给药途径"].extend(dosingWay)
                                # print(f"给药频率1 = ", end="")
                                # print(node['药嘱.给药频率'])
                                node["病人资料.年龄"][sentence_split].update(
                                    {"药嘱.用药间隔周期": node["药嘱.用药间隔周期"], "药嘱.给药频率": node["药嘱.给药频率"],
                                     "药嘱.使用疗程": node["药嘱.使用疗程"], "药嘱.每次给药剂量": node["药嘱.每次给药剂量"],
                                     "药嘱.每公斤每次剂量": node["药嘱.每公斤每次剂量"], "药嘱.每公斤每次极量": node["药嘱.每公斤每次极量"],
                                     "药嘱.每天剂量": node["药嘱.每天剂量"], "药嘱.给药目的": node["药嘱.给药目的"],
                                     "药嘱.每公斤每日剂量": node["药嘱.每公斤每日剂量"], "药嘱.每公斤每日极量": node["药嘱.每公斤每日极量"],
                                     "药嘱.次极量": node["药嘱.次极量"], "药嘱.日极量": node["药嘱.日极量"], "药嘱.每公斤总剂量": node["药嘱.每公斤总剂量"],
                                     "药嘱.每公斤每分钟滴速": node["药嘱.每公斤每分钟滴速"], "药嘱.滴速": node["药嘱.滴速"],
                                     "药嘱.持续时间": node["药嘱.持续时间"], "药嘱.每平方米每次剂量": node["药嘱.每平方米每次剂量"],
                                     "药嘱.每平方米每次极量": node["药嘱.每平方米每次极量"], "药嘱.每平方米每日剂量": node["药嘱.每平方米每日剂量"],
                                     "药嘱.每平方米每日极量": node["药嘱.每平方米每日极量"], "药嘱.每平方米总剂量": node["药嘱.每平方米总剂量"],
                                     "药嘱.总剂量": node["药嘱.总剂量"], "药嘱.药品剂型": node["药嘱.药品剂型"],
                                     "药嘱.给药时机": node["药嘱.给药时机"]})
                                # print('time used: {} sec'.format((datetime.now() - start).total_seconds()))

                            # 对reader0，即前面未被分割的进行提取
                            sentence_0_token = tokenizer.tokenize(reader0)
                            # print(sentence_token)
                            # print('your input is:{}'.format(sentence_token))
                            input_ids, input_mask, segment_ids, label_ids = convert(sentence_0_token)

                            feed_dict = {input_ids_p: input_ids,
                                         input_mask_p: input_mask}
                            # run session get current feed_dict result
                            pred_ids_result = sess.run([pred_ids], feed_dict)
                            pred_label_result = convert_id_to_label(pred_ids_result, id2label)
                            # 输出预测标签
                            # print(pred_label_result[0])

                            node["药嘱.用药间隔周期"] = ""
                            node["药嘱.给药频率"] = []
                            node["药嘱.使用疗程"] = ""
                            node["药嘱.每次给药剂量"] = ""
                            node["药嘱.每公斤每次剂量"] = ""
                            node["药嘱.每公斤每次极量"] = ""
                            node["药嘱.每天剂量"] = ""
                            node["药嘱.给药目的"] = ""
                            node["药嘱.每公斤每日剂量"] = ""

                            node["药嘱.每公斤每日极量"] = ""
                            node["药嘱.次极量"] = ""
                            node["药嘱.日极量"] = ""
                            node["药嘱.每公斤总剂量"] = ""
                            node["药嘱.每公斤每分钟滴速"] = ""
                            node["药嘱.滴速"] = ""
                            node["药嘱.持续时间"] = ""
                            node["药嘱.每平方米每次剂量"] = ""
                            node["药嘱.每平方米每次极量"] = ""
                            node["药嘱.每平方米每日剂量"] = ""
                            node["药嘱.每平方米每日极量"] = ""
                            node["药嘱.每平方米总剂量"] = ""
                            node["药嘱.总剂量"] = ""
                            node["药嘱.药品剂型"] = ""
                            node["药嘱.给药时机"] = ""

                            node["药嘱.用药间隔周期"], node["药嘱.给药频率"], node["药嘱.使用疗程"], \
                            node["药嘱.每次给药剂量"], node["药嘱.每公斤每次剂量"], node["药嘱.每公斤每次极量"], \
                            node["药嘱.每天剂量"], node["药嘱.给药目的"], node["药嘱.每公斤每日剂量"], \
                            node["药嘱.每公斤每日极量"], node["药嘱.次极量"], node["药嘱.日极量"], \
                            node["药嘱.每公斤总剂量"], node["药嘱.每公斤每分钟滴速"], node["药嘱.滴速"], \
                            node["药嘱.持续时间"], node["药嘱.每平方米每次剂量"], node["药嘱.每平方米每次极量"], \
                            node["药嘱.每平方米每日剂量"], node["药嘱.每平方米每日极量"], node["药嘱.每平方米总剂量"], \
                            node["药嘱.总剂量"], node["药嘱.药品剂型"], node["药嘱.给药时机"], dosingWay = search_field(sentence_0_token,
                                                                                                       pred_label_result[
                                                                                                           0],
                                                                                                       others_field)
                            node["药嘱.给药途径"].extend(dosingWay)

                            # 去重
                            node["药嘱.给药途径"] = list(set(node["药嘱.给药途径"]))
                            dosingWay = ""
                            for i, tmp in enumerate(node["药嘱.给药途径"]):
                                if i == 0:
                                    dosingWay += tmp
                                else:
                                    dosingWay += f"|{tmp}"
                            dosingWay = re.sub("(或|、)", "|", dosingWay)
                            dosingWay = dosingWay.replace("肌肉注射", "肌内注射")
                            node["药嘱.给药途径"] = dosingWay

                            # print(node['药嘱.给药频率'])
                            print('具体字段识别耗费时间: {} sec'.format((datetime.now() - start).total_seconds()))
                            if else_sign:
                                node["病人资料.年龄"]["else"] = {}
                                node["病人资料.年龄"]["else"].update(
                                    {"药嘱.用药间隔周期": node["药嘱.用药间隔周期"], "药嘱.给药频率": node["药嘱.给药频率"],
                                     "药嘱.使用疗程": node["药嘱.使用疗程"], "药嘱.每次给药剂量": node["药嘱.每次给药剂量"],
                                     "药嘱.每公斤每次剂量": node["药嘱.每公斤每次剂量"], "药嘱.每公斤每次极量": node["药嘱.每公斤每次极量"],
                                     "药嘱.每天剂量": node["药嘱.每天剂量"], "药嘱.给药目的": node["药嘱.给药目的"],
                                     "药嘱.每公斤每日剂量": node["药嘱.每公斤每日剂量"], "药嘱.每公斤每日极量": node["药嘱.每公斤每日极量"],
                                     "药嘱.次极量": node["药嘱.次极量"], "药嘱.日极量": node["药嘱.日极量"], "药嘱.每公斤总剂量": node["药嘱.每公斤总剂量"],
                                     "药嘱.每公斤每分钟滴速": node["药嘱.每公斤每分钟滴速"], "药嘱.滴速": node["药嘱.滴速"],
                                     "药嘱.持续时间": node["药嘱.持续时间"], "药嘱.每平方米每次剂量": node["药嘱.每平方米每次剂量"],
                                     "药嘱.每平方米每次极量": node["药嘱.每平方米每次极量"], "药嘱.每平方米每日剂量": node["药嘱.每平方米每日剂量"],
                                     "药嘱.每平方米每日极量": node["药嘱.每平方米每日极量"], "药嘱.每平方米总剂量": node["药嘱.每平方米总剂量"],
                                     "药嘱.总剂量": node["药嘱.总剂量"], "药嘱.药品剂型": node["药嘱.药品剂型"],
                                     "药嘱.给药时机": node["药嘱.给药时机"]})
                                NODE['用法用量'].update({"病人资料.年龄": node["病人资料.年龄"]})
                            else:
                                NODE['用法用量'].update({"药嘱.用药间隔周期": node["药嘱.用药间隔周期"], "药嘱.给药频率": node["药嘱.给药频率"],
                                                     "药嘱.使用疗程": node["药嘱.使用疗程"], "药嘱.每次给药剂量": node["药嘱.每次给药剂量"],
                                                     "药嘱.每公斤每次剂量": node["药嘱.每公斤每次剂量"], "药嘱.每公斤每次极量": node["药嘱.每公斤每次极量"],
                                                     "药嘱.每天剂量": node["药嘱.每天剂量"], "药嘱.给药目的": node["药嘱.给药目的"],
                                                     "药嘱.每公斤每日剂量": node["药嘱.每公斤每日剂量"], "药嘱.每公斤每日极量": node["药嘱.每公斤每日极量"],
                                                     "药嘱.次极量": node["药嘱.次极量"], "药嘱.日极量": node["药嘱.日极量"],
                                                     "药嘱.每公斤总剂量": node["药嘱.每公斤总剂量"],
                                                     "药嘱.每公斤每分钟滴速": node["药嘱.每公斤每分钟滴速"], "药嘱.滴速": node["药嘱.滴速"],
                                                     "药嘱.持续时间": node["药嘱.持续时间"], "药嘱.每平方米每次剂量": node["药嘱.每平方米每次剂量"],
                                                     "药嘱.每平方米每次极量": node["药嘱.每平方米每次极量"],
                                                     "药嘱.每平方米每日剂量": node["药嘱.每平方米每日剂量"],
                                                     "药嘱.每平方米每日极量": node["药嘱.每平方米每日极量"],
                                                     "药嘱.每平方米总剂量": node["药嘱.每平方米总剂量"],
                                                     "药嘱.总剂量": node["药嘱.总剂量"], "药嘱.药品剂型": node["药嘱.药品剂型"],
                                                     "药嘱.给药时机": node["药嘱.给药时机"], "病人资料.年龄": node["病人资料.年龄"]})
                            NODE["给药途径"].update({"药嘱.给药途径 属于": node["药嘱.给药途径"]})
                            field_dict.update(others_field)
                elif dd == '注意事项.txt' or dd == '禁忌.txt' or dd == '禁忌症.txt':
                    alter_sign = True
                    with open(f"{files[0]}/{dd}", 'r', encoding='utf-8') as txt:
                        while True:
                            reader = txt.readline()
                            if reader == "":
                                break
                            reader = reader.replace("\n", "")
                            # 生化指标.肝功能检查
                            mm_21_1 = re.finditer(compiled_rule_21_1, reader)
                            # 生化指标.肾功能检查
                            mm_22_1 = re.finditer(compiled_rule_22_1, reader)

                            # 生化指标.肝/肾功能检查
                            node["生化指标.肝功能检查.谷丙转氨酶"], node["生化指标.肾功能检查.内生肌酐清除率"] = search_functionTest(mm_21_1, mm_22_1)
                            if node["生化指标.肝功能检查.谷丙转氨酶"]:
                                NODE["生化指标"].update({"生化指标.肝功能检查.谷丙转氨酶": node["生化指标.肝功能检查.谷丙转氨酶"]})
                            if node["生化指标.肾功能检查.内生肌酐清除率"]:
                                NODE["生化指标"].update({"生化指标.肾功能检查.内生肌酐清除率": node["生化指标.肾功能检查.内生肌酐清除率"]})

                            if reader.find("过敏") != -1:
                                node["病人资料.过敏药品列表"] = "属于 "
                            if node["病人资料.过敏药品列表"]:
                                NODE["过敏"].update({"病人资料.过敏药品列表": node["病人资料.过敏药品列表"]})
                # 如果用法用量内没有年龄，再看接种对象里有没有年龄标记
                elif dd == '接种对象.txt':
                    alter_sign = True
                    with open(f"{files[0]}/{dd}", 'r', encoding='utf-8') as txt:
                        while True:
                            reader = txt.readline()
                            if reader == "":
                                break
                            reader = reader.replace("\n", "")
                            if reader.find("男") != -1:
                                node["病人资料.性别"] = "属于 男"
                                # print(node["sex"][0])
                            elif reader.find("女") != -1:
                                node["病人资料.性别"] = "属于 女"
                                # print(node["病人资料.性别"][0])
                            if node["病人资料.性别"]:
                                NODE["用法用量"].update({"病人资料.性别": node["病人资料.性别"]})
                elif dd == '儿童用药.txt':
                    alter_sign = True
                    with open(f"{files[0]}/{dd}", 'r', encoding='utf-8') as txt:
                        while True:
                            reader = txt.readline()
                            if reader == "":
                                break
                            reader = reader.replace("\n", "")
                            if reader.find("新生儿") != -1 or reader.find("早产儿") != -1:
                                if reader.find("禁用") != -1:
                                    # node["病人资料.年龄"].update({"< 8 天": {"用药建议": "禁用"}})
                                    node["病人资料.年龄"].update({"< 8 天": ""})
                                elif reader.find("慎用") != -1 or reader.find("谨慎") != -1:
                                    # node["病人资料.年龄"].update({"< 8 天": {"用药建议": "慎用"}})
                                    node["病人资料.年龄"].update({"< 8 天": ""})
                                # print(node["病人资料.年龄"][0])
                            elif reader.find("婴幼儿") != -1 or reader.find("婴儿") != -1 or reader.find("幼儿") != -1:
                                if reader.find("禁用") != -1:
                                    # node["病人资料.年龄"].update({">=且< 28 天 1 岁": {"用药建议": "禁用"}})
                                    node["病人资料.年龄"].update({">= 28天 且< 1岁": ""})
                                elif reader.find("慎用") != -1 or reader.find("谨慎") != -1:
                                    # node["病人资料.年龄"].update({">=且< 28 天 1 岁": {"用药建议": "慎用"}})
                                    node["病人资料.年龄"].update({">= 28天 且< 1岁": {}})
                                # print(node["病人资料.年龄"][0])
                            elif reader.find("儿童") != -1:
                                if reader.find("禁用") != -1:
                                    # node["病人资料.年龄"].update({"<= 14 岁": {"用药建议": "禁用"}})
                                    node["病人资料.年龄"].update({"<= 14 岁": {}})
                                elif reader.find("慎用") != -1 or reader.find("谨慎") != -1:
                                    # node["病人资料.年龄"].update({"<= 14 岁": {"用药建议": "慎用"}})
                                    node["病人资料.年龄"].update({"<= 14 岁": {}})
                                # print(node["病人资料.年龄"][0])
                            elif reader.find("禁用") != -1:
                                # node["病人资料.年龄"].update({"<= 14 岁": {"用药建议": "禁用"}})
                                node["病人资料.年龄"].update({"<= 14 岁": {}})
                            elif reader.find("慎用") != -1 or reader.find("谨慎") != -1:
                                # node["病人资料.年龄"].update({"<= 14 岁": {"用药建议": "慎用"}})
                                node["病人资料.年龄"].update({"<= 14 岁": {}})
                            if node["病人资料.年龄"]:
                                NODE["用法用量"].update({"病人资料.年龄": node["病人资料.年龄"]})
                elif dd == '老年用药.txt':
                    alter_sign = True
                    with open(f"{files[0]}/{dd}", 'r', encoding='utf-8') as txt:
                        while True:
                            reader = txt.readline()
                            if reader == "":
                                break
                            # node["病人资料.年龄"] = []
                            reader = reader.replace("\n", "")
                            if reader.find("老年人") != -1:
                                # node["病人资料.年龄"].update({">= 60 岁": {"用药建议": "慎用"}})
                                node["病人资料.年龄"].update({">= 60 岁": {}})
                                # print(node["age"][0])
                            elif reader.find("高龄") != -1:
                                # node["病人资料.年龄"].update({">= 80 岁": {"用药建议": "慎用"}})
                                node["病人资料.年龄"].update({">= 80 岁": {}})
                                # print(node["age"][0])
                            elif reader.find("禁用") != -1:
                                # node["病人资料.年龄"].update({">= 60 岁": {"用药建议": "禁用"}})
                                node["病人资料.年龄"].update({">= 60 岁": {}})
                            elif reader.find("慎用") != -1 or reader.find("谨慎") != -1:
                                # node["病人资料.年龄"].update({">= 60 岁": {"用药建议": "慎用"}})
                                node["病人资料.年龄"].update({">= 60 岁": {}})
                            if node["病人资料.年龄"]:
                                NODE["用法用量"].update({"病人资料.年龄": node["病人资料.年龄"]})
                elif dd == '孕妇及哺乳期妇女用药.txt':
                    alter_sign = True
                    with open(f"{files[0]}/{dd}", 'r', encoding='utf-8') as txt:
                        while True:
                            reader = txt.readline()
                            if reader == "":
                                break
                            reader = reader.replace("\n", "")
                            if reader.find("孕妇") != -1:
                                node["病人资料.是否怀孕"] = "= True"
                                # node["pregnancy"].append(f"病人资料.是否怀孕 = True 药品建议：孕妇{}")
                                # if reader.find("禁用") != -1:
                                #     node["病人资料.是否怀孕"] = "= True"  # 用药建议：禁用
                                # elif reader.find("慎用") != -1 or reader.find("谨慎") != -1 or reader.find("慎重") != -1:
                                #     node["病人资料.是否怀孕"] = "= True"  # 用药建议：慎用
                            if reader.find("哺乳") != -1:
                                node["病人资料.是否哺乳"] = "= True"
                                # if reader.find("禁用") != -1:
                                #     node["病人资料.是否哺乳"] = "= True"  # 用药建议：禁用
                                # elif reader.find("慎用") != -1 or reader.find("谨慎") != -1 or reader.find("慎重") != -1:
                                #     node["病人资料.是否哺乳"] = "= True"  # 用药建议：慎用
                            if node["病人资料.是否怀孕"]:
                                NODE["孕产"].update({"病人资料.是否怀孕": node["病人资料.是否怀孕"]})
                            if node["病人资料.是否哺乳"]:
                                NODE["孕产"].update({"病人资料.是否哺乳": node["病人资料.是否哺乳"]})

            if alter_sign == True:
                drug_json = jsonFormatting(drug_json)
                # 去重
                for key in field_dict:
                    field_dict[key] = list(set(field_dict[key]))

                print('总识别耗费时间: {} sec'.format((datetime.now() - start_all).total_seconds()))
                # print(f"field_dict = {field_dict}")
                return field_dict, drug_json


def convert_id_to_label(pred_ids_result, idx2label):
    """
    将id形式的结果转化为真实序列结果
    :param pred_ids_result:
    :param idx2label:
    :return:
    """
    result = []
    for row in range(batch_size):
        curr_seq = []
        for ids in pred_ids_result[row][0]:
            if ids == 0:
                break
            curr_label = idx2label[ids]
            if curr_label in ['[CLS]', '[SEP]']:
                continue
            curr_seq.append(curr_label)
        result.append(curr_seq)
    return result


def strage_combined_link_org_loc(tokens, tags):
    """
    组合策略
    :param pred_label_result:
    :param types:
    :return:
    """

    def print_output(data, type):
        line = []
        line.append(type)
        for i in data:
            line.append(i.word)
        print(', '.join(line))

    params = None
    eval = Result(params)
    if len(tokens) > len(tags):
        tokens = tokens[:len(tags)]
    medicationInterva, dosingFrequence, useTreatment, eachTimeDose, doseKGEachTime, maximumDoseKGEachTime, dayDose, \
    dosingPurpose, doseKGEachDay, maximumDoseKGEachDay, maximumeachTimeDose, maximumDayDose, doseKGTotal, \
    dropRateKGEachMinute, dropRate, duration, doseM2EachTime, maximumDoseM2EachTime, doseM2EachDay, maximumDoseM2EachDay, \
    doseM2Total, doseTotal, dosingDosageForm, dosingTime, dosingWay, age = eval.get_result(tokens, tags)

    print_output(medicationInterva, '药嘱.间隔周期')
    print_output(dosingFrequence, '药嘱.给药频率')
    print_output(useTreatment, '药嘱.使用疗程')
    print_output(eachTimeDose, '药嘱.每次给药剂量')
    print_output(doseKGEachTime, '药嘱.每公斤每次剂量')
    print_output(maximumDoseKGEachTime, '药嘱.每公斤每次极量')
    print_output(dayDose, '药嘱.每天剂量')
    print_output(dosingPurpose, '药嘱.给药目的')
    print_output(doseKGEachDay, '药嘱.每公斤每日剂量')

    print_output(maximumDoseKGEachDay, '药嘱.每公斤每日极量')
    print_output(maximumeachTimeDose, '药嘱.次极量')
    print_output(maximumDayDose, '药嘱.日极量')
    print_output(doseKGTotal, '药嘱.每公斤总剂量')
    print_output(dropRateKGEachMinute, '药嘱.每公斤每分钟滴速')
    print_output(dropRate, '药嘱.滴速')
    print_output(duration, '药嘱.持续时间')
    print_output(doseM2EachTime, '药嘱.每平方米每次剂量')
    print_output(maximumDoseM2EachTime, '药嘱.每平方米每次极量')
    print_output(doseM2EachDay, '药嘱.每平方米每日剂量')
    print_output(maximumDoseM2EachDay, '药嘱.每平方米每日极量')
    print_output(doseM2Total, '药嘱.每平方米总剂量')
    print_output(doseTotal, '药嘱.总剂量')
    print_output(dosingDosageForm, '药嘱.药品剂型')
    print_output(dosingTime, '药嘱.给药时机')
    print_output(dosingWay, '药嘱.给药途径 属于')
    print_output(age, '年龄')


def convert_single_example(ex_index, example, label_list, max_seq_length, tokenizer, mode):
    """
    将一个样本进行分析，然后将字转化为id, 标签转化为id,然后结构化到InputFeatures对象中
    :param ex_index: index
    :param example: 一个样本
    :param label_list: 标签列表
    :param max_seq_length:
    :param tokenizer:
    :param mode:
    :return:
    """
    label_map = {}
    # 1表示从1开始对label进行index化
    for (i, label) in enumerate(label_list, 1):
        label_map[label] = i
    # 保存label->index 的map
    if not os.path.exists(os.path.join(model_dir, 'label2id.pkl')):
        with codecs.open(os.path.join(model_dir, 'label2id.pkl'), 'wb') as w:
            pickle.dump(label_map, w)

    tokens = example
    # tokens = tokenizer.tokenize(example.text)
    # 序列截断
    if len(tokens) >= max_seq_length - 1:
        tokens = tokens[0:(max_seq_length - 2)]  # -2 的原因是因为序列需要加一个句首和句尾标志
    ntokens = []
    segment_ids = []
    label_ids = []
    ntokens.append("[CLS]")  # 句子开始设置CLS 标志
    segment_ids.append(0)
    # append("O") or append("[CLS]") not sure!
    label_ids.append(label_map["[CLS]"])  # O OR CLS 没有任何影响，不过我觉得O 会减少标签个数,不过拒收和句尾使用不同的标志来标注，使用LCS 也没毛病
    for i, token in enumerate(tokens):
        ntokens.append(token)
        segment_ids.append(0)
        label_ids.append(0)
    ntokens.append("[SEP]")  # 句尾添加[SEP] 标志
    segment_ids.append(0)
    # append("O") or append("[SEP]") not sure!
    label_ids.append(label_map["[SEP]"])
    input_ids = tokenizer.convert_tokens_to_ids(ntokens)  # 将序列中的字(ntokens)转化为ID形式
    input_mask = [1] * len(input_ids)

    # padding, 使用
    while len(input_ids) < max_seq_length:
        input_ids.append(0)
        input_mask.append(0)
        segment_ids.append(0)
        # we don't concerned about it!
        label_ids.append(0)
        ntokens.append("**NULL**")
        # label_mask.append(0)
    # print(len(input_ids))
    assert len(input_ids) == max_seq_length
    assert len(input_mask) == max_seq_length
    assert len(segment_ids) == max_seq_length
    assert len(label_ids) == max_seq_length
    # assert len(label_mask) == max_seq_length

    # 结构化为一个类
    feature = InputFeatures(
        input_ids=input_ids,
        input_mask=input_mask,
        segment_ids=segment_ids,
        label_ids=label_ids,
        # label_mask = label_mask
    )
    return feature


class Pair(object):
    def __init__(self, word, start, end, type, merge=False):
        self.__word = word
        self.__start = start
        self.__end = end
        self.__merge = merge
        self.__types = type

    @property
    def start(self):
        return self.__start

    @property
    def end(self):
        return self.__end

    @property
    def merge(self):
        return self.__merge

    @property
    def word(self):
        return self.__word

    @property
    def types(self):
        return self.__types

    @word.setter
    def word(self, word):
        self.__word = word

    @start.setter
    def start(self, start):
        self.__start = start

    @end.setter
    def end(self, end):
        self.__end = end

    @merge.setter
    def merge(self, merge):
        self.__merge = merge

    @types.setter
    def types(self, type):
        self.__types = type

    def __str__(self) -> str:
        line = []
        line.append('entity:{}'.format(self.__word))
        line.append('start:{}'.format(self.__start))
        line.append('end:{}'.format(self.__end))
        line.append('merge:{}'.format(self.__merge))
        line.append('types:{}'.format(self.__types))
        return '\t'.join(line)


class Result(object):
    def __init__(self, config):
        self.config = config
        self.medicationInterva = []
        self.dosingFrequence = []
        self.useTreatment = []
        self.eachTimeDose = []
        self.doseKGEachTime = []
        self.maximumDoseKGEachTime = []
        self.dayDose = []
        self.dosingPurpose = []
        self.doseKGEachDay = []

        self.maximumDoseKGEachDay = []
        self.maximumeachTimeDose = []
        self.maximumDayDose = []
        self.doseKGTotal = []
        self.dropRateKGEachMinute = []
        self.dropRate = []
        self.duration = []
        self.doseM2EachTime = []
        self.maximumDoseM2EachTime = []
        self.doseM2EachDay = []
        self.maximumDoseM2EachDay = []
        self.doseM2Total = []
        self.doseTotal = []
        self.dosingDosageForm = []
        self.dosingTime = []
        self.dosingWay = []
        self.age = []

        self.others = []

    def get_result(self, tokens, tags, config=None):
        # 先获取标注结果
        self.result_to_json(tokens, tags)
        return self.medicationInterva, self.dosingFrequence, self.useTreatment, self.eachTimeDose, self.doseKGEachTime, \
               self.maximumDoseKGEachTime, self.dayDose, self.dosingPurpose, self.doseKGEachDay, \
               self.maximumDoseKGEachDay, self.maximumeachTimeDose, self.maximumDayDose, self.doseKGTotal, \
               self.dropRateKGEachMinute, self.dropRate, self.duration, self.doseM2EachTime, self.maximumDoseM2EachTime, \
               self.doseM2EachDay, self.maximumDoseM2EachDay, self.doseM2Total, self.doseTotal, self.dosingDosageForm, \
               self.dosingTime, self.dosingWay, self.age

    def result_to_json(self, string, tags):
        """
        将模型标注序列和输入序列结合 转化为结果
        :param string: 输入序列
        :param tags: 标注结果
        :return:
        """
        item = {"entities": []}
        entity_name = ""
        entity_start = 0
        idx = 0
        last_tag = ''

        for char, tag in zip(string, tags):
            if tag[0] == "S":
                self.append(char, idx, idx + 1, tag[2:])
                item["entities"].append({"word": char, "start": idx, "end": idx + 1, "type": tag[2:]})
            elif tag[0] == "B":
                if entity_name != '':
                    self.append(entity_name, entity_start, idx, last_tag[2:])
                    item["entities"].append(
                        {"word": entity_name, "start": entity_start, "end": idx, "type": last_tag[2:]})
                    entity_name = ""
                entity_name += char
                entity_start = idx
            elif tag[0] == "I":
                entity_name += char
            elif tag[0] == "O":
                if entity_name != '':
                    self.append(entity_name, entity_start, idx, last_tag[2:])
                    item["entities"].append(
                        {"word": entity_name, "start": entity_start, "end": idx, "type": last_tag[2:]})
                    entity_name = ""
            else:
                entity_name = ""
                entity_start = idx
            idx += 1
            last_tag = tag
        if entity_name != '':
            self.append(entity_name, entity_start, idx, last_tag[2:])
            item["entities"].append({"word": entity_name, "start": entity_start, "end": idx, "type": last_tag[2:]})
        return item

    def append(self, word, start, end, tag):
        if tag == 'medicationInterva':
            self.medicationInterva.append(Pair(word, start, end, 'medicationInterva'))
        elif tag == 'dosingFrequence':
            self.dosingFrequence.append(Pair(word, start, end, 'dosingFrequence'))
        elif tag == 'useTreatment':
            self.useTreatment.append(Pair(word, start, end, 'useTreatment'))
        elif tag == 'eachTimeDose':
            self.eachTimeDose.append(Pair(word, start, end, 'eachTimeDose'))
        elif tag == 'doseKGEachTime':
            self.doseKGEachTime.append(Pair(word, start, end, 'doseKGEachTime'))
        elif tag == 'maximumDoseKGEachTime':
            self.maximumDoseKGEachTime.append(Pair(word, start, end, 'maximumDoseKGEachTime'))
        elif tag == 'dayDose':
            self.dayDose.append(Pair(word, start, end, 'dayDose'))
        elif tag == 'dosingPurpose':
            self.dosingPurpose.append(Pair(word, start, end, 'dosingPurpose'))
        elif tag == 'doseKGEachDay':
            self.doseKGEachDay.append(Pair(word, start, end, 'doseKGEachDay'))
        elif tag == 'maximumDoseKGEachDay':
            self.maximumDoseKGEachDay.append(Pair(word, start, end, 'maximumDoseKGEachDay'))
        elif tag == 'maximumeachTimeDose':
            self.maximumeachTimeDose.append(Pair(word, start, end, 'maximumeachTimeDose'))
        elif tag == 'maximumDayDose':
            self.maximumDayDose.append(Pair(word, start, end, 'maximumDayDose'))
        elif tag == 'doseKGTotal':
            self.doseKGTotal.append(Pair(word, start, end, 'doseKGTotal'))
        elif tag == 'dropRateKGEachMinute':
            self.dropRateKGEachMinute.append(Pair(word, start, end, 'dropRateKGEachMinute'))
        elif tag == 'dropRate':
            self.dropRate.append(Pair(word, start, end, 'dropRate'))
        elif tag == 'duration':
            self.duration.append(Pair(word, start, end, 'duration'))
        elif tag == 'doseM2EachTime':
            self.doseM2EachTime.append(Pair(word, start, end, 'doseM2EachTime'))
        elif tag == 'maximumDoseM2EachTime':
            self.maximumDoseM2EachTime.append(Pair(word, start, end, 'maximumDoseM2EachTime'))
        elif tag == 'doseM2EachDay':
            self.doseM2EachDay.append(Pair(word, start, end, 'doseM2EachDay'))
        elif tag == 'maximumDoseM2EachDay':
            self.maximumDoseM2EachDay.append(Pair(word, start, end, 'maximumDoseM2EachDay'))
        elif tag == 'doseM2Total':
            self.doseM2Total.append(Pair(word, start, end, 'doseM2Total'))
        elif tag == 'doseTotal':
            self.doseTotal.append(Pair(word, start, end, 'doseTotal'))
        elif tag == 'dosingDosageForm':
            self.dosingDosageForm.append(Pair(word, start, end, 'dosingDosageForm'))
        elif tag == 'dosingTime':
            self.dosingTime.append(Pair(word, start, end, 'dosingTime'))
        elif tag == 'dosingWay':
            self.dosingWay.append(Pair(word, start, end, 'dosingWay'))
        elif tag == 'age':
            self.age.append(Pair(word, start, end, 'age'))
        else:
            self.others.append(Pair(word, start, end, tag))


def search_age(sentence, tokens, tags):
    def output_list(data):
        line = []
        for i in data:
            line.append(i.word)
        return line

    params = None
    eval = Result(params)
    if len(tokens) > len(tags):
        tokens = tokens[:len(tags)]
    # print(f"tags_age = {tags}")
    medicationInterva, dosingFrequence, useTreatment, eachTimeDose, doseKGEachTime, maximumDoseKGEachTime, dayDose, \
    dosingPurpose, doseKGEachDay, maximumDoseKGEachDay, maximumeachTimeDose, maximumDayDose, doseKGTotal, \
    dropRateKGEachMinute, dropRate, duration, doseM2EachTime, maximumDoseM2EachTime, doseM2EachDay, \
    maximumDoseM2EachDay, doseM2Total, doseTotal, dosingDosageForm, dosingTime, dosingWay, age = eval.get_result(tokens,
                                                                                                                 tags)

    age = output_list(age)
    age_field = age
    # print(f"age = {age}")
    # 病人资料.年龄
    PATTERN_23_1 = r"([\u4e00-\u9fa5]*)(\d+\.?\d*)(个月|月|岁)?(～|\-|~|—|±|一|或者|及|或|至)?(\d+\.?\d*)?(个月|月|岁)([\u4e00-\u9fa5]*)"
    PATTERN_23_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5])(个月|月|岁)?(～|\-|~|—|±|一|或者|及|或|至)([\u4e00-\u9fa5])(个月|月|岁)([\u4e00-\u9fa5]*)"
    PATTERN_23_3 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5])(个月|月|岁)([\u4e00-\u9fa5]*)"

    # 病人资料.年龄
    compiled_rule_23_1 = re.compile(PATTERN_23_1)
    compiled_rule_23_2 = re.compile(PATTERN_23_2)
    compiled_rule_23_3 = re.compile(PATTERN_23_3)

    split_str = ""
    # 病人资料.年龄
    age_split_record = {}
    for tmp in age:
        mm_23_1 = re.search(compiled_rule_23_1, tmp)
        mm_23_2 = re.search(compiled_rule_23_2, tmp)
        mm_23_3 = re.search(compiled_rule_23_3, tmp)

        # PATTERN_23_1 = r"([\u4e00-\u9fa5]*)(\d+\.?\d*)(个月|月|岁)?(～|\-|~|—|±|一|或者|及|或|至)?(\d+\.?\d*)?(个月|月|岁)([\u4e00-\u9fa5]*)"
        sign = True
        a_list = []
        if mm_23_1:
            sign = False
            # print(f"mm_23_1 = {mm_23_1.group(0)}")
            str_s = mm_23_1.group(1)
            m1 = mm_23_1.group(2)
            symbol = mm_23_1.group(4)
            m2 = mm_23_1.group(5)
            m3 = mm_23_1.group(3)
            m4 = mm_23_1.group(6)
            str_e = mm_23_1.group(7)
            if m3 == "个月":
                m3 = "月"
            if m4 == "个月":
                m4 = "月"
            if not sign:
                # a_list.append("病人资料.年龄")
                if str_e.find("以上") != -1 or str_s.find("大于") != -1 or str_s.find("超过") != -1:
                    a_list.append(">")
                    a_list.append(m1)
                    if m3 != None:
                        a_list.append(m3)
                    else:
                        a_list.append(m4)
                elif str_e.find("以内") != -1 or str_e.find("以下") != -1 or str_s.find("小于") != -1 or str_s.find(
                        "少于") != -1:
                    a_list.append("<")
                    if m2 != None:
                        if symbol == "±":
                            a_list.append(float(m1) + float(m2))
                        else:
                            a_list.append(m2)
                    else:
                        a_list.append(m1)
                    a_list.append(m4)
                elif symbol == "±":
                    a_list.append(float(m1) - float(m2))
                    if m3 != None:
                        a_list.append(m3)
                    else:
                        a_list.append(m4)
                    a_list.append(float(m1) + float(m2))
                    a_list.append(m4)
                elif m2 != None:
                    a_list.append(">=")
                    a_list.append(m1)
                    if m3 != None:
                        a_list.append(m3)
                    else:
                        a_list.append(m4)
                    a_list.append("且<=")
                    a_list.append(m2)
                    a_list.append(m4)
                else:
                    a_list.append("=")
                    a_list.append(m1)
                    a_list.append(m4)
                buff = ''
                for temp in a_list:
                    buff += str(temp)
                    buff += ' '
                age = mm_23_1.group(0)
                age_split_record.update({age: buff})
        # PATTERN_23_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5])(个月|月|岁)?(～|\-|~|—|±|一|或者|及|或|至)([\u4e00-\u9fa5])(个月|月|岁)([\u4e00-\u9fa5]*)"
        if sign and mm_23_2:
            # print(f"mm_23_2 = {mm_23_2.group(0)}")
            sign = False
            str_s = mm_23_2.group(1)
            m1 = mm_23_2.group(2)
            symbol = mm_23_2.group(4)
            m2 = mm_23_2.group(5)
            m3 = mm_23_2.group(3)
            m4 = mm_23_2.group(6)
            str_e = mm_23_2.group(7)
            if m1 == "两" or m1 == "俩":
                m1 = 2
            if m2 == "两" or m2 == "俩":
                m2 = 2
            if m3 == "个月":
                m3 = "月"
            if m4 == "个月":
                m4 = "月"
            if not is_number(m1):
                sign = True
            elif is_number(m1) == 2:
                # 一 二 三 ---> 1 2 3
                m1 = round(unicodedata.numeric(m1))
            if not is_number(m2):
                sign = True
            elif is_number(m2) == 2:
                # 一 二 三 ---> 1 2 3
                m2 = round(unicodedata.numeric(m2))
            if not sign:
                if str_e.find("以上") != -1 or str_s.find("大于") != -1 or str_s.find("超过") != -1:
                    a_list.append(">")
                    a_list.append(m1)
                    a_list.append(m3)
                elif str_e.find("以内") != -1 or str_e.find("以下") != -1 or str_s.find("小于") != -1 or str_s.find(
                        "少于") != -1:
                    a_list.append("<")
                    if symbol == "±":
                        a_list.append(float(m1) + float(m2))
                    else:
                        a_list.append(m2)
                    a_list.append(m4)
                elif symbol == "±":
                    a_list.append(float(m1) - float(m2))
                    a_list.append(m3)
                    a_list.append(float(m1) + float(m2))
                    a_list.append(m4)
                else:
                    a_list.append(">=")
                    a_list.append(m1)
                    if m3:
                        a_list.append(m3)
                    else:
                        a_list.append(m4)
                    a_list.append("且<=")
                    a_list.append(m2)
                    a_list.append(m4)
                buff = ''
                for temp in a_list:
                    # print(temp, end=' ')
                    buff += str(temp)
                    buff += ' '
                age = mm_23_2.group(0)
                age_split_record.update({age: buff})
        # PATTERN_23_3 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5])(个月|月|岁)([\u4e00-\u9fa5]*)"
        if sign and mm_23_3:
            # print(f"mm_23_3 = {mm_23_3.group(0)}")
            sign = False
            str_s = mm_23_3.group(1)
            m1 = mm_23_3.group(2)
            m3 = mm_23_3.group(3)
            str_e = mm_23_3.group(4)
            if m1 == "两" or m1 == "俩":
                m1 = 2
            if not is_number(m1):
                sign = True
            elif is_number(m1) == 2:
                m1 = round(unicodedata.numeric(m1))
            if not sign:
                if str_e.find("以上") != -1 or str_s.find("大于") != -1 or str_s.find("超过") != -1:
                    a_list.append(">")
                elif str_e.find("以内") != -1 or str_e.find("以下") != -1 or str_s.find("小于") != -1 or str_s.find(
                        "少于") != -1:
                    a_list.append("<")
                else:
                    a_list.append("=")
                a_list.append(m1)
                a_list.append(m3)
                buff = ''
                for temp in a_list:
                    buff += str(temp)
                    buff += ' '
                age = mm_23_3.group(0)
                age_split_record.update({age: buff})
        for key in age_split_record:
            split_str = split_str + "|" + key

    # print(f"age_split_record = {age_split_record}")
    buff = re.split(f"(新生儿|婴儿|儿童|幼儿|老人|老年人|高龄|早产儿|成人{split_str})", sentence)
    phase_dict_original = {}
    for i, key in enumerate(buff):
        if i == 0:
            continue
        elif i % 2 == 1:
            if key in phase_dict_original.keys():
                phase_dict_original[key] += buff[i + 1]
            else:
                phase_dict_original.update({key: buff[i + 1]})
        else:
            pass
    # print(f"phase_dict_original = {phase_dict_original}")
    # print(f"len1 = {len(phase_dict_original)}")
    phase_dict_final = {}
    for key in phase_dict_original:
        if key in age_split_record.keys():
            # print(age_split_record[key])
            value = age_split_record[key]
            phase_dict_final[value] = phase_dict_original[key]
        elif key == "新生儿":
            age_field.append("新生儿")
            phase_dict_final["< 28天"] = phase_dict_original["新生儿"]
        elif key == "婴儿":
            age_field.append("婴儿")
            phase_dict_final[">= 28天 且< 1岁"] = phase_dict_original["婴儿"]
        elif key == "儿童":
            age_field.append("儿童")
            phase_dict_final[">= 2月 且< 2岁"] = phase_dict_original["儿童"]
        elif key == "幼儿":
            age_field.append("幼儿")
            phase_dict_final[">= 1岁 且< 3岁"] = phase_dict_original["幼儿"]
        elif key == "老人":
            age_field.append("老人")
            phase_dict_final[">= 60岁"] = phase_dict_original["老人"]
        elif key == "老年人":
            age_field.append("老年人")
            phase_dict_final[">= 60岁"] = phase_dict_original["老年人"]
        elif key == "高龄":
            age_field.append("高龄")
            phase_dict_final[">= 80岁"] = phase_dict_original["高龄"]
        elif key == "早产儿":
            age_field.append("早产儿")
            phase_dict_final["< 8天"] = phase_dict_original["早产儿"]
        elif key == "成人":
            age_field.append("成人")
            phase_dict_final[">= 14岁 且< 60岁"] = phase_dict_original["成人"]
        else:
            value = phase_dict_original[key]
            phase_dict_final[value] = phase_dict_original[key]

    # print(f"phase_dict_final = {phase_dict_final}")
    return phase_dict_final, buff[0], age_field


def search_field(tokens, tags, others_field):
    def output_list(data):
        line = []
        for i in data:
            line.append(i.word)
        return line

    params = None
    eval = Result(params)
    if len(tokens) > len(tags):
        tokens = tokens[:len(tags)]

    medicationInterva, dosingFrequence, useTreatment, eachTimeDose, doseKGEachTime, maximumDoseKGEachTime, dayDose, \
    dosingPurpose, doseKGEachDay, maximumDoseKGEachDay, maximumeachTimeDose, maximumDayDose, doseKGTotal, \
    dropRateKGEachMinute, dropRate, duration, doseM2EachTime, maximumDoseM2EachTime, doseM2EachDay, maximumDoseM2EachDay, \
    doseM2Total, doseTotal, dosingDosageForm, dosingTime, dosingWay, age = eval.get_result(tokens, tags)

    # print(f"tags_field = {tags}")

    reader_medicationInterva = output_list(medicationInterva)
    reader_dosingFrequence = output_list(dosingFrequence)
    reader_useTreatment = output_list(useTreatment)
    reader_eachTimeDose = output_list(eachTimeDose)
    reader_doseKGEachTime = output_list(doseKGEachTime)
    reader_maximumDoseKGEachTime = output_list(maximumDoseKGEachTime)
    reader_dayDose = output_list(dayDose)
    reader_dosingPurpose = output_list(dosingPurpose)
    reader_doseKGEachDay = output_list(doseKGEachDay)

    reader_maximumDoseKGEachDay = output_list(maximumDoseKGEachDay)
    reader_maximumeachTimeDose = output_list(maximumeachTimeDose)
    reader_maximumDayDose = output_list(maximumDayDose)
    reader_doseKGTotal = output_list(doseKGTotal)
    reader_dropRateKGEachMinute = output_list(dropRateKGEachMinute)
    reader_dropRate = output_list(dropRate)
    reader_duration = output_list(duration)
    reader_doseM2EachTime = output_list(doseM2EachTime)
    reader_maximumDoseM2EachTime = output_list(maximumDoseM2EachTime)
    reader_doseM2EachDay = output_list(doseM2EachDay)
    reader_maximumDoseM2EachDay = output_list(maximumDoseM2EachDay)
    reader_doseM2Total = output_list(doseM2Total)
    reader_doseTotal = output_list(doseTotal)
    reader_dosingDosageForm = output_list(dosingDosageForm)
    reader_dosingTime = output_list(dosingTime)
    reader_dosingWay = output_list(dosingWay)

    others_field["用药间隔周期"].extend(reader_medicationInterva)
    others_field["给药频率"].extend(reader_dosingFrequence)
    others_field["使用疗程"].extend(reader_useTreatment)
    others_field["每次给药剂量"].extend(reader_eachTimeDose)
    others_field["每公斤每次极量"].extend(reader_doseKGEachTime)
    others_field["每天剂量"].extend(reader_dayDose)
    others_field["给药目的"].extend(reader_dosingPurpose)
    others_field["每公斤每日剂量"].extend(reader_doseKGEachDay)
    others_field["每公斤每日极量"].extend(reader_maximumDoseKGEachDay)
    others_field["次极量"].extend(reader_maximumeachTimeDose)
    others_field["日极量"].extend(reader_maximumDayDose)
    others_field["每公斤总剂量"].extend(reader_doseKGTotal)
    others_field["每公斤每分钟滴速"].extend(reader_dropRateKGEachMinute)
    others_field["滴速"].extend(reader_dropRate)
    others_field["持续时间"].extend(reader_duration)
    others_field["每平方米每次剂量"].extend(reader_doseM2EachTime)
    others_field["每平方米每次极量"].extend(reader_maximumDoseM2EachTime)
    others_field["每平方米每日剂量"].extend(reader_doseM2EachDay)
    others_field["每平方米每日极量"].extend(reader_maximumDoseM2EachDay)
    others_field["每平方米总剂量"].extend(reader_doseM2Total)
    others_field["总剂量"].extend(reader_doseTotal)
    others_field["药品剂型"].extend(reader_dosingDosageForm)
    others_field["给药时机"].extend(reader_dosingTime)
    others_field["给药途径"].extend(reader_dosingWay)

    # 药嘱.给药频率
    # 匹配类型：一日一次或隔日一次
    PATTERN_1_1 = r"([\u4e00-\u9fa5]|\d+)(周|日|天|小时|时)([\u4e00-\u9fa5]|\d+)(次)[～|\-|~|\－|—|±|一|或|至]([\u4e00-\u9fa5]|\d+)(周|日|天|小时|时)([\u4e00-\u9fa5]|\d+)(次)"
    # 匹配类型：一周1次，一周1或2次，12小时给|用药一次，每周静脉给药一次，每周进行1-2次，24小时更换一次，24小时服用一次
    PATTERN_1_2 = r"([\u4e00-\u9fa5]|\d+)(周|日|天|小时|时)更?换?进?行?静?脉?[给|用|服]?用?药?([\u4e00-\u9fa5][～|\-|~|\－|—|±|一|或|至][\u4e00-\u9fa5]|[\u4e00-\u9fa5]|\d+[～|\-|~|—|一|－|或|至]\d+|\d+)(次)"  # |片|支|坎|包
    # 匹配类型：一周或两周1次，一或两周1次
    PATTERN_1_3 = r"([\u4e00-\u9fa5]|\d+)(周|日|天|小时|时)?[～|\-|~|\－|—|±|一|或|至]([\u4e00-\u9fa5]|\d+)(周|日|天|小时|时)换?药?([\u4e00-\u9fa5]|\d+)(次)"
    # 匹配类型：每日早晚各一次
    PATTERN_1_4 = r"([\u4e00-\u9fa5]|\d+)(日|天)早晚各?服?([\u4e00-\u9fa5]|\d+)(次)"
    # 匹配类型：早晚各一次，早晚一次
    PATTERN_1_5 = r"早晚各?服?([\u4e00-\u9fa5]|\d+)次"
    # 匹配类型：3次/日
    PATTERN_1_6 = r"([\u4e00-\u9fa5]|\d+)次/日"
    # 匹配类型：每晚睡前一次，每晚一次
    PATTERN_1_7 = r"每晚睡?前?([\u4e00-\u9fa5]|\d+)次"
    # 匹配类型：分早中两次服用，分早、中二次服用，分2~4次服用，分2～4次，分2次，三至四次
    PATTERN_1_8 = r"分早?、?中?([\u4e00-\u9fa5][～|\-|~|\－|—|±|一|或|至]?[\u4e00-\u9fa5]?|\d+[～|\-|~|\－|—|±|一|或|至]\d+|\d+)次"
    # 匹配类型：顿服，单剂，单次，分次
    PATTERN_1_9 = r"([\u4e00-\u9fa5]*)(顿服|单剂|单次|分次)([\u4e00-\u9fa5]*)"

    # 药嘱.每天剂量    每次给药剂量
    # 匹配类型：超过1片~2片/以上
    # 匹配类型：0.1~3g，0.1g~3g
    PATTERN_2_1 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+\.?\d*)(喷|丸|吸|盒|颗|张|U|揿|万U|l|复方|条|粒|MIU|掀|kg|mg|IU|ul|ku|滴|包|豪居里|μg|g|枚|支|袋|片|瓶|听|贴|mCi|万IU|ml|kcal|BU)?(～|\-|~|\－|—|±|一|或者|或|至)([\u4e00-\u9fa5]|\d+\.?\d*)(喷|丸|吸|盒|颗|张|U|揿|万U|l|复方|条|粒|MIU|掀|kg|mg|IU|ul|ku|滴|包|豪居里|μg|g|枚|支|袋|片|瓶|听|贴|mCi|万IU|ml|kcal|BU)([\u4e00-\u9fa5]*)"
    # 匹配类型：2g
    PATTERN_2_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+\.?\d*)(喷|丸|吸|盒|颗|张|U|揿|万U|l|复方|条|粒|MIU|掀|kg|mg|IU|ul|ku|滴|包|豪居里|μg|g|枚|支|袋|片|瓶|听|贴|mCi|万IU|ml|kcal|BU)([\u4e00-\u9fa5]*)"

    # 药嘱.使用疗程
    # 匹配类型：3-4周
    PATTERN_4_1 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(次|日|个月|月|时|天|年|周)?(～|\-|~|\－|—|±|一|或者|或|至)([\u4e00-\u9fa5]|\d+)(次|日|个月|月|时|天|年|周)([\u4e00-\u9fa5]*)"
    # 匹配类型：6周
    PATTERN_4_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(次|日|个月|月|时|天|年|周)([\u4e00-\u9fa5]*)"

    # 药嘱.每公斤每日剂量     药嘱.每公斤每次剂量     药嘱.每公斤每次极量
    #  匹配类型：日15～25mg/kg;1.5g/kg体重/天
    PATTERN_5_1 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+\.?\d*)(喷|丸|吸|盒|颗|张|U|揿|万U|l|复方|条|粒|MIU|掀|kg|mg|IU|ul|ku|滴|包|豪居里|μg|g|枚|支|袋|片|瓶|听|贴|mCi|万IU|ml|kcal|BU)?(～|\-|~|\－|—|±|一|或者|或|至)([\u4e00-\u9fa5]|\d+\.?\d*)(喷|丸|吸|盒|颗|张|U|揿|万U|l|复方|条|粒|MIU|掀|kg|mg|IU|ul|ku|滴|包|豪居里|μg|g|枚|支|袋|片|瓶|听|贴|mCi|万IU|ml|kcal|BU)\/?k?g?m?2?([\u4e00-\u9fa5]*)"
    PATTERN_5_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+\.?\d*)(喷|丸|吸|盒|颗|张|U|揿|万U|l|复方|条|粒|MIU|掀|kg|mg|IU|ul|ku|滴|包|豪居里|μg|g|枚|支|袋|片|瓶|听|贴|mCi|万IU|ml|kcal|BU)\/?k?g?m?2?([\u4e00-\u9fa5]*)"

    # 药嘱.用药间隔周期
    #  匹配类型：间隔4-6周；
    PATTERN_10_1 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(～|\-|~|\－|—|±|一|或者|或|至)([\u4e00-\u9fa5]|\d+)(个月|日|年|秒|天|小时|月|分|周|时)([\u4e00-\u9fa5]*)"
    #  匹配类型：间隔四周
    PATTERN_10_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(个月|日|年|秒|天|小时|月|分|周|时)([\u4e00-\u9fa5]*)"

    # 药嘱.持续时间
    #  匹配类型：滴注时间依据剂量不同至少为60分钟或90分钟以上。M10065769
    PATTERN_9_1 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(小时|分|分钟|时|秒|日|天|个月|月)?(～|\-|~|\－|—|±|一|或者|或|至)([\u4e00-\u9fa5]|\d+)(小时|分|分钟|时|秒|日|天|个月|月)([\u4e00-\u9fa5]*)"
    # 匹配类型： 滴1小时以上M10066654;滴注不宜超过6小时;缓慢给药（1分钟以上 ;注射时间不少于3分钟
    PATTERN_9_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(小时|分|分钟|时|秒|日|天|个月|月)([\u4e00-\u9fa5]*)"

    dosingTime_str = r"(进食时|早上|早饭后|早餐前和晚餐中|下午|午餐中和晚餐前|晚餐前和睡前|三餐前和晚餐后|早餐前和午餐中|早餐中和午餐后|晨间和晚餐后|餐中|午餐中和晚餐中|晨间|早餐前和午餐后|晨间和上午|顿服|晨间和晚餐前|午餐前和睡前|午餐前|晨间和下午|晨间和午餐中|经前|午餐后和睡前|晚餐中|早餐中和睡前|晚餐后和睡前|上午|早餐后和睡前|睡前|早餐中和午餐前|午餐前和晚餐前|午餐后和晚餐后|早餐后和晚餐前|早餐中|午餐中和晚餐后|术后|晨间和午餐后|早餐中和晚餐前|餐后|早餐后和午餐后|三餐前和睡前|晨间和早餐中|\
        |晚餐后|下午和晚上|空腹|早餐前和晚餐前|早餐前|晚餐前|早餐后和晚餐后|术前|经期|上午和晚上|术中|午餐后和晚餐中|上午和中午|上午和下午|午餐前和晚餐后|晨间和中午|早餐后和午餐中|经后|午餐后和晚餐前|中午和下午|晨间和午餐前|早餐中和晚餐后|早餐后|早餐前和晚餐后|中午|三餐中和睡前|餐前|早餐前和睡前|午餐中和睡前|晚上|三餐后和睡前|早餐中和晚餐中|晨间和睡前|必要时|晚餐中和睡前|早餐中和午餐中|早餐后和晚餐中|午餐后|午餐中|晨间和晚上|\
        |中午和晚上|晨间和早餐后|早餐后和午餐前|早餐前和午餐前|晨间和晚餐中|午餐前和晚餐中)"
    # 药嘱.给药时机
    PATTERN_30_1 = r"([\u4e00-\u9fa5]*)" + dosingTime_str + r"([\u4e00-\u9fa5]*)"
    PATTERN_30_2 = r"([\u4e00-\u9fa5]*)" + dosingTime_str + r"[或|和]" + dosingTime_str + r"([\u4e00-\u9fa5]*)"

    # 药嘱.给药目的
    PATTERN_20_1 = r"([\u4e00-\u9fa5]*)(预防|治疗|应急)([\u4e00-\u9fa5]*)"

    # 药嘱.给药频率
    compiled_rule_1_1 = re.compile(PATTERN_1_1)
    compiled_rule_1_2 = re.compile(PATTERN_1_2)
    compiled_rule_1_3 = re.compile(PATTERN_1_3)
    compiled_rule_1_4 = re.compile(PATTERN_1_4)
    compiled_rule_1_5 = re.compile(PATTERN_1_5)
    compiled_rule_1_6 = re.compile(PATTERN_1_6)
    compiled_rule_1_7 = re.compile(PATTERN_1_7)
    compiled_rule_1_8 = re.compile(PATTERN_1_8)
    compiled_rule_1_9 = re.compile(PATTERN_1_9)
    # 药嘱.每天剂量    # 药嘱.每次给药剂量
    compiled_rule_2_1 = re.compile(PATTERN_2_1)
    compiled_rule_2_2 = re.compile(PATTERN_2_2)

    # 药嘱.使用疗程
    compiled_rule_4_1 = re.compile(PATTERN_4_1)
    compiled_rule_4_2 = re.compile(PATTERN_4_2)

    # 药嘱.每公斤每日剂量  药嘱.每公斤每次剂量   药嘱.每公斤每次极量
    compiled_rule_5_1 = re.compile(PATTERN_5_1)
    compiled_rule_5_2 = re.compile(PATTERN_5_2)

    # 药嘱.用药间隔周期
    compiled_rule_10_1 = re.compile(PATTERN_10_1)
    compiled_rule_10_2 = re.compile(PATTERN_10_2)

    # 药嘱.持续时间
    compiled_rule_9_1 = re.compile(PATTERN_9_1)
    compiled_rule_9_2 = re.compile(PATTERN_9_2)

    # 药嘱.给药目的
    compiled_rule_20_1 = re.compile(PATTERN_20_1)
    # 药嘱.给药时机
    compiled_rule_30_1 = re.compile(PATTERN_30_1)
    compiled_rule_30_2 = re.compile(PATTERN_30_2)

    medicationInterva = None
    dosingFrequence = []
    useTreatment = None
    eachTimeDose = None
    doseKGEachTime = None
    maximumDoseKGEachTime = None
    dayDose = None
    dosingPurpose = None
    doseKGEachDay = None

    maximumDoseKGEachDay = None
    maximumeachTimeDose = None
    maximumDayDose = None
    doseKGTotal = None
    dropRateKGEachMinute = None
    dropRate = None
    duration = None
    doseM2EachTime = None
    maximumDoseM2EachTime = None
    doseM2EachDay = None
    maximumDoseM2EachDay = None
    doseM2Total = None
    doseTotal = None
    dosingDosageForm = ""
    dosingTime = None
    dosingWay = ""

    # 药嘱.用药间隔周期
    # print(f"用药间隔周期 = {reader_medicationInterva}")
    for tmp in reader_medicationInterva:
        mm_10_1 = re.search(compiled_rule_10_1, tmp)
        mm_10_2 = re.search(compiled_rule_10_2, tmp)
        buf = search_medicationInterva(mm_10_1, mm_10_2)
        if buf:
            medicationInterva = buf
            break

    # 药嘱.给药频率
    for tmp in reader_dosingFrequence:
        mm_1_1 = re.search(compiled_rule_1_1, tmp)
        mm_1_2 = re.search(compiled_rule_1_2, tmp)
        mm_1_3 = re.search(compiled_rule_1_3, tmp)
        mm_1_4 = re.search(compiled_rule_1_4, tmp)
        mm_1_5 = re.search(compiled_rule_1_5, tmp)
        mm_1_6 = re.search(compiled_rule_1_6, tmp)
        mm_1_7 = re.search(compiled_rule_1_7, tmp)
        mm_1_8 = re.search(compiled_rule_1_8, tmp)
        mm_1_9 = re.search(compiled_rule_1_9, tmp)
        buf = search_dosingFrequence(mm_1_1, mm_1_2, mm_1_3, mm_1_4, mm_1_5, mm_1_6, mm_1_7, mm_1_8, mm_1_9)
        if buf:
            dosingFrequence.append(buf)
    # 给药频率 去重
    dosingFrequence = list(set(dosingFrequence))

    # 药嘱.使用疗程
    # print(f"使用疗程 = {reader_useTreatment}")
    buf = []
    dose_list = []
    for tmp in reader_useTreatment:
        mm_4_1 = re.search(compiled_rule_4_1, tmp)
        mm_4_2 = re.search(compiled_rule_4_2, tmp)
        buf_tmp, aa, bb = search_useTreatment(mm_4_1, mm_4_2)
        if buf_tmp:
            buf.append(buf_tmp)
        if aa:
            dose_list.append(aa)
        if bb:
            dose_list.append(bb)
    # print(buf)
    # print(dose_list)
    if len(buf):
        if len(dose_list) > 1:
            for i, aa in enumerate(dose_list):
                if aa.find(".") == -1:
                    dose_list[i] = int(aa)
                else:
                    dose_list[i] = float(aa)
            min_1 = min(dose_list)
            max_1 = max(dose_list)
            str_list = buf[0].split()
            if len(str_list) > 3:
                buf[0] = buf[0].replace(str_list[1], str(min_1))
                buf[0] = buf[0].replace(str_list[4], str(max_1))
            else:
                buf[0] = buf[0].replace(str_list[0], ">=")
                buf[0] = buf[0].replace(str_list[1], str(min_1))
                buf[0] += " <= "
                buf[0] += str(max_1)
                buf[0] += " "
                buf[0] += str_list[2]
        useTreatment = buf[0]

    buf = []
    dose_list = []
    # 药嘱.每次给药剂量
    # print(f"每次给药剂量 = {reader_eachTimeDose}")
    for tmp in reader_eachTimeDose:
        first_use = False
        mm_3_1 = re.search(compiled_rule_2_1, tmp)
        mm_3_2 = re.search(compiled_rule_2_2, tmp)
        buf_tmp, aa, bb = search_dayDose_eachTimeDose(mm_3_1, mm_3_2)
        if buf_tmp:
            buf.append(buf_tmp)
        # print(f"min_1 = {min_1}， max_1 = {max_1}")
        # 判断是否是首剂使用
        if tmp.find("初始剂量") != -1 or tmp.find("首剂") != -1:
            first_use = True
            break
        if aa:
            dose_list.append(aa)
        if bb:
            dose_list.append(bb)
    if len(buf):
        if first_use == True:
            eachTimeDose = "药嘱.首剂使用 = true" + buf[0]
        else:
            if len(dose_list) > 1:
                for i, aa in enumerate(dose_list):
                    if aa.find(".") == -1:
                        dose_list[i] = int(aa)
                    else:
                        dose_list[i] = float(aa)
                min_1 = min(dose_list)
                max_1 = max(dose_list)
                str_list = buf[0].split()
                if len(str_list) > 3:
                    buf[0] = buf[0].replace(str_list[1], str(min_1))
                    buf[0] = buf[0].replace(str_list[4], str(max_1))
                else:
                    buf[0] = buf[0].replace(str_list[0], ">=")
                    buf[0] = buf[0].replace(str_list[1], str(min_1))
                    buf[0] += " <= "
                    buf[0] += str(max_1)
                    buf[0] += " "
                    buf[0] += str_list[2]
            eachTimeDose = buf[0]

    # 药嘱.每公斤每次剂量
    # print(f"每公斤每次剂量 = {reader_doseKGEachTime}")
    for tmp in reader_doseKGEachTime:
        mm_6_1 = re.search(compiled_rule_5_1, tmp)
        mm_6_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_doseKGEachDay_doseKGEachTime(mm_6_1, mm_6_2)
        if buf:
            doseKGEachTime = buf
            break

    # 药嘱.每公斤每次极量
    # print(f"每公斤每次极量 = {reader_maximumDoseKGEachTime}")
    for tmp in reader_maximumDoseKGEachTime:
        mm_7_1 = re.search(compiled_rule_5_1, tmp)
        mm_7_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_maximumDoseKGEachTime(mm_7_1, mm_7_2)
        if buf:
            maximumDoseKGEachTime = buf
            break

    #  药嘱.每天剂量
    # print(f"每天剂量 = {reader_dayDose}")
    buf = []
    dose_list = []
    for tmp in reader_dayDose:
        mm_2_1 = re.search(compiled_rule_2_1, tmp)
        mm_2_2 = re.search(compiled_rule_2_2, tmp)
        buf_tmp, aa, bb = search_dayDose_eachTimeDose(mm_2_1, mm_2_2)
        if buf_tmp:
            buf.append(buf_tmp)
        if aa:
            dose_list.append(aa)
        if bb:
            dose_list.append(bb)
    if len(buf):
        if len(dose_list) > 1:
            for i, aa in enumerate(dose_list):
                if aa.find(".") == -1:
                    dose_list[i] = int(aa)
                else:
                    dose_list[i] = float(aa)
            min_1 = min(dose_list)
            max_1 = max(dose_list)
            str_list = buf[0].split()
            if len(str_list) > 3:
                buf[0] = buf[0].replace(str_list[1], str(min_1))
                buf[0] = buf[0].replace(str_list[4], str(max_1))
            else:
                buf[0] = buf[0].replace(str_list[0], ">=")
                buf[0] = buf[0].replace(str_list[1], str(min_1))
                buf[0] += " <= "
                buf[0] += str(max_1)
                buf[0] += " "
                buf[0] += str_list[2]
        dayDose = buf[0]

    # 药嘱.给药目的
    # print(f"给药目的 = {reader_dosingPurpose}")
    for tmp in reader_dosingPurpose:
        mm_20_1 = re.search(compiled_rule_20_1, tmp)
        buf = search_dosingPurpose(mm_20_1)
        if buf:
            dosingPurpose = buf
            break

    # 药嘱.每公斤每日剂量
    # print(f"每公斤每日剂量 = {reader_doseKGEachDay}")
    for tmp in reader_doseKGEachDay:
        mm_5_1 = re.search(compiled_rule_5_1, tmp)
        mm_5_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_doseKGEachDay_doseKGEachTime(mm_5_1, mm_5_2)
        if buf:
            doseKGEachDay = buf
            break

    # 药嘱.每公斤每日极量
    # print(f"每公斤每日极量 = {reader_maximumDoseKGEachDay}")
    for tmp in reader_maximumDoseKGEachDay:
        mm_11_1 = re.search(compiled_rule_5_1, tmp)
        mm_11_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_maximumDoseKGEachTime(mm_11_1, mm_11_2)
        if buf:
            maximumDoseKGEachDay = buf
            break

    # 药嘱.次极量
    # print(f"次极量 = {reader_maximumeachTimeDose}")
    for tmp in reader_maximumeachTimeDose:
        mm_11_1 = re.search(compiled_rule_5_1, tmp)
        mm_11_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_maximumDoseKGEachTime(mm_11_1, mm_11_2)
        if buf:
            maximumeachTimeDose = buf
            break

    # 药嘱.日极量
    # print(f"日极量 = {reader_maximumDayDose}")
    for tmp in reader_maximumDayDose:
        mm_11_1 = re.search(compiled_rule_5_1, tmp)
        mm_11_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_maximumDoseKGEachTime(mm_11_1, mm_11_2)
        if buf:
            maximumDayDose = buf
            break

    # 药嘱.每公斤总剂量
    # print(f"每公斤总剂量 = {reader_doseKGTotal}")
    for tmp in reader_doseKGTotal:
        mm_11_1 = re.search(compiled_rule_5_1, tmp)
        mm_11_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_maximumDoseKGEachTime(mm_11_1, mm_11_2)
        if buf:
            doseKGTotal = buf
            break

    # 药嘱.每公斤每分钟滴速
    # print(f"每公斤每分钟滴速 = {reader_dropRateKGEachMinute}")
    for tmp in reader_dropRateKGEachMinute:
        mm_6_1 = re.search(compiled_rule_5_1, tmp)
        mm_6_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_doseKGEachDay_doseKGEachTime(mm_6_1, mm_6_2)
        if buf:
            dropRateKGEachMinute = buf
            break

    # 药嘱.滴速
    # print(f"滴速 = {reader_dropRate}")
    for tmp in reader_dropRate:
        mm_6_1 = re.search(compiled_rule_5_1, tmp)
        mm_6_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_doseKGEachDay_doseKGEachTime(mm_6_1, mm_6_2)
        if buf:
            dropRate = buf
            break

    # 药嘱.持续时间
    # print(f"持续时间 = {reader_duration}")
    for tmp in reader_duration:
        mm_9_1 = re.search(compiled_rule_9_1, tmp)
        mm_9_2 = re.search(compiled_rule_9_2, tmp)
        buf = search_duration(mm_9_1, mm_9_2)
        if buf:
            duration = buf
            break

    # 药嘱.每平方米每次剂量
    # print(f"每平方米每次剂量 = {reader_doseM2EachTime}")
    for tmp in reader_doseM2EachTime:
        mm_6_1 = re.search(compiled_rule_5_1, tmp)
        mm_6_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_doseKGEachDay_doseKGEachTime(mm_6_1, mm_6_2)
        if buf:
            doseM2EachTime = buf
            break

    # 药嘱.每平方米每次极量
    # print(f"每平方米每次极量 = {reader_maximumDoseM2EachTime}")
    for tmp in reader_maximumDoseM2EachTime:
        mm_7_1 = re.search(compiled_rule_5_1, tmp)
        mm_7_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_maximumDoseKGEachTime(mm_7_1, mm_7_2)
        if buf:
            maximumDoseM2EachTime = buf
            break

    # 药嘱.每平方米每日剂量
    # print(f"每平方米每日剂量 = {reader_doseM2EachDay}")
    for tmp in reader_doseM2EachDay:
        mm_6_1 = re.search(compiled_rule_5_1, tmp)
        mm_6_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_doseKGEachDay_doseKGEachTime(mm_6_1, mm_6_2)
        if buf:
            doseM2EachDay = buf
            break

    # 药嘱.每平方米每日极量
    # print(f"每平方米每日极量 = {reader_maximumDoseM2EachDay}")
    for tmp in reader_maximumDoseM2EachDay:
        mm_7_1 = re.search(compiled_rule_5_1, tmp)
        mm_7_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_maximumDoseKGEachTime(mm_7_1, mm_7_2)
        if buf:
            maximumDoseM2EachDay = buf
            break

    # 药嘱.每平方米总剂量
    # print(f"每平方米总剂量 = {reader_doseM2Total}")
    for tmp in reader_doseM2Total:
        mm_11_1 = re.search(compiled_rule_5_1, tmp)
        mm_11_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_maximumDoseKGEachTime(mm_11_1, mm_11_2)
        if buf:
            doseM2Total = buf
            break

    # 药嘱.总剂量
    # print(f"总剂量 = {reader_doseTotal}")
    for tmp in reader_doseTotal:
        mm_11_1 = re.search(compiled_rule_5_1, tmp)
        mm_11_2 = re.search(compiled_rule_5_2, tmp)
        buf = search_maximumDoseKGEachTime(mm_11_1, mm_11_2)
        if buf:
            doseTotal = buf
            break

    # 药嘱.药品剂型
    # 去重
    reader_dosingDosageForm = list(set(reader_dosingDosageForm))
    for i, tmp in enumerate(reader_dosingDosageForm):
        if i == 0:
            dosingDosageForm += tmp
        else:
            dosingDosageForm += f"|{tmp}"
    dosingDosageForm = re.sub("(或|、)", "|", dosingDosageForm)
    # print(f"药品剂型 = {dosingDosageForm}")

    # 药嘱.给药时机
    # print(f"给药时机 = {reader_dosingTime}")
    for tmp in reader_dosingTime:
        mm_30_1 = re.search(compiled_rule_30_1, tmp)
        mm_30_2 = re.search(compiled_rule_30_2, tmp)
        buf = search_dosingTime(mm_30_1, mm_30_2)
        if buf:
            dosingTime = buf
            break

    # print(f"给药途径 = {dosingWay}")
    return medicationInterva, dosingFrequence, useTreatment, eachTimeDose, doseKGEachTime, maximumDoseKGEachTime, \
           dayDose, dosingPurpose, doseKGEachDay, maximumDoseKGEachDay, maximumeachTimeDose, maximumDayDose, doseKGTotal, \
           dropRateKGEachMinute, dropRate, duration, doseM2EachTime, maximumDoseM2EachTime, doseM2EachDay, \
           maximumDoseM2EachDay, doseM2Total, doseTotal, dosingDosageForm, dosingTime, reader_dosingWay


# 创建所需文件夹
def create_directory():
    if (not os.path.exists(f"{absp}/data/instructions")):
        print("create directory instructions...")
        os.mkdir(f"{absp}/data/instructions")
    if (not os.path.exists(f"{absp}/data/instructions_split")):
        print("create directory instructions_split...")
        os.mkdir(f"{absp}/data/instructions_split")

    if (not os.path.exists(f"{absp}/data/instructions_docx")):
        print("create directory instructions_docx...")
        os.mkdir(f"{absp}/data/instructions_docx")


# doc和docx文档归类
def exchange_folder_file_doc_docx():
    print("exchange folder file doc docx...")
    path_doc = f"{absp}/data/instructions_doc"
    path_docx = f"{absp}/data/instructions_docx"
    path_instructions = f"{absp}/data/instructions"
    for files in os.listdir(path_instructions):
        file_name = os.path.splitext(files)[0]
        file_property = os.path.splitext(files)[1]
        if file_property == '.doc':
            path1 = f"{path_instructions}/{files}"
            path2 = f"{path_doc}/{files}"
            shutil.move(path1, path2)
        elif file_property == '.docx':
            path1 = f"{path_instructions}/{files}"
            path2 = f"{path_docx}/{files}"
            shutil.move(path1, path2)


def read_docx(files):
    create_directory()
    file_name = os.path.splitext(files)[0]
    file_property = os.path.splitext(files)[1]
    content_docx = ""
    path_instructions = f"{absp}/data/instructions"
    if file_property == ".doc":
        content_docx = "<h3>请选择.docx文件<br><b>注意：</b><br>1.该docx文件不能通过.doc文件直接改后缀名得到.docx。<br>2.应该由.doc文件另存为.docx得到。或者新建docx，并将内容填入</h3>"
    if not content_docx:
        try:
            path = f'{path_instructions}/{file_name}.docx'
            document = Document(path)
            # 读取文档中所有的段落列表
            ps = document.paragraphs
            # 每个段落有两个属性：style和text
            ps_detail = [(x.text, x.style.name) for x in ps]
            content_docx = ""
            for line in ps_detail:
                content_docx = content_docx + line[0] + '<br>'
        except:
            content_docx = "<center><h2>解决方法一：关闭word<br></h1></center>"
    return content_docx




# 将单个Docx文本分割成txt
def DocxSplitTxt(file_name):
    # print("split docx to txt...")
    path_instructions = f"{absp}/data/instructions"
    if not os.path.isfile(f'{path_instructions}/{file_name}.docx'):
        path_instructions = f"{absp}/data/instructions_docx"
    PATTERN_1 = r'【(.*)】([\u4e00-\u9fa5]*)'
    PATTERN_2 = r'\[(.*)\]([\u4e00-\u9fa5]*)'
    compiled_rule_1 = re.compile(PATTERN_1)
    compiled_rule_2 = re.compile(PATTERN_2)
    if (not os.path.exists(f"{absp}/data/instructions_split/{file_name}")):
        os.mkdir(f"{absp}/data/instructions_split/{file_name}")
    document = Document(f'{path_instructions}/{file_name}.docx')
    # 读取文档中所有的段落列表
    ps = document.paragraphs
    # 每个段落有两个属性：style和text
    ps_detail = [(x.text, x.style.name) for x in ps]
    first_txt = False
    find_name = False
    head = None
    find_sign = False
    for line in ps_detail:
        temp = ''.join(line[0].split())
        temp.strip()
        if (temp == '' or temp == '【】'):
            continue
        mm_1 = re.match(compiled_rule_1, temp)
        mm_2 = re.match(compiled_rule_2, temp)
        if mm_1:
            if mm_1.group(1) == "药品名称":
                find_name = True
            else:
                find_name = False
            # 标题：种/族，转换为：种 族
            head = mm_1.group(1).replace("/", "\\")
            file = open(f"{absp}/data/instructions_split/{file_name}/" + head + ".txt", 'w', encoding='utf-8')
            if temp.split("】")[1]:
                for i, buf in enumerate(temp.split("】")):
                    if i == 0:
                        continue
                    file.write(buf)
                first_txt = False
            else:
                first_txt = True
            file.close()
        elif mm_2:
            if mm_2.group(1) == "药品名称":
                find_name = True
            else:
                find_name = False
            # 标题：种/族，转换为：种 族
            head = mm_2.group(1).replace("/", "\\")
            file = open(f"{absp}/data/instructions_split/{file_name}/" + head + ".txt", 'w', encoding='utf-8')
            if temp.split("]")[1]:
                for i, buf in enumerate(temp.split("]")):
                    if i == 0:
                        continue
                    file.write(buf)
                first_txt = False
            else:
                first_txt = True
            file.close()
        else:
            if first_txt == False:
                continue  # 去掉可划分前的所有段落
            if find_name == True:
                buf = line[0]
                with open(f"{absp}/data/instructions_split/{file_name}/" + head + ".txt", 'a', encoding='utf-8') as f1:
                    f1.write(buf)
                f1.close()
                if buf.find("通用名称") != -1:
                    buf = buf.replace(":", "")
                    buf = buf.replace("：", "")
                    drug_name = buf.split("通用名称")[1]
                    with open(f"{absp}/data/instructions_split/{file_name}/首行药品名称.txt", 'w', encoding='utf-8') as f1:
                        f1.write(drug_name)
                    f1.close()
                    find_sign = True
                else:
                    if not find_sign:
                        with open(f"{absp}/data/instructions_split/{file_name}/首行药品名称.txt", 'w',
                                  encoding='utf-8') as f1:
                            f1.write(buf)
                        f1.close()
                        find_sign = True
            else:
                with open(f"{absp}/data/instructions_split/{file_name}/" + head + ".txt", 'a', encoding='utf-8') as f1:
                    f1.write(line[0])
                f1.close()
    if not find_sign:
        with open(f"{absp}/data/instructions_split/{file_name}/首行药品名称.txt", 'w', encoding='utf-8') as f1:
            f1.write(file_name)
        f1.close()


# 数字转中文
def _to_chinese4(num):
    _MAPPING = (
        u'零', u'一', u'二', u'三', u'四', u'五', u'六', u'七', u'八', u'九', u'十', u'十一', u'十二', u'十三', u'十四', u'十五', u'十六',
        u'十七',
        u'十八', u'十九')
    _P0 = (u'', u'十', u'百', u'千',)
    _S4 = 10 ** 4
    assert (0 <= num and num < _S4)
    if num < 20:
        return _MAPPING[num]
    else:
        lst = []
        while num >= 10:
            lst.append(num % 10)
            num = num / 10
        lst.append(num)
        c = len(lst)  # 位数
        result = u''
        for idx, val in enumerate(lst):
            val = int(val)
            if val != 0:
                result += _P0[idx] + _MAPPING[val]
                if idx < c - 1 and lst[idx + 1] == 0:
                    result += u'零'
        return result[::-1]


def is_number(s):
    try:  # 如果能运行float(s)语句，返回True（字符串s是浮点数）
        s = float(s)
        return 1
    except ValueError:  # ValueError为Python的一种标准异常，表示"传入无效的参数"
        pass  # 如果引发了ValueError这种异常，不做任何事情（pass：不做任何事情，一般用做占位语句）
    try:
        s = unicodedata.numeric(s)  # 把一个表示数字的字符串转换为浮点数返回的函数
        return 2
    except (TypeError, ValueError):
        pass
    return 0


# 药嘱.药品名称
def search_drugName(file_path):
    with open(f"{file_path}/首行药品名称.txt", 'r', encoding='utf-8') as txt:
        dm_list = []
        while True:
            reader = txt.readline()
            reader = reader.replace(" ", "")
            if reader == "":
                break
            reader = re.split(r"药品名称：", reader)
            # dm_list.append("药嘱.药品名称")
            if reader[0] == '':
                dm_list.append(reader[1])
            else:
                dm_list.append(reader[0])
            return dm_list[0]
    return '无'


# 药嘱.给药频率
def search_dosingFrequence(mm_1_1, mm_1_2, mm_1_3, mm_1_4, mm_1_5, mm_1_6, mm_1_7, mm_1_8, mm_1_9):
    sign = True
    df_list = []
    if mm_1_1:
        # 匹配类型：一天1次或隔天1次
        sign = False
        m1 = mm_1_1.group(1)
        m2 = mm_1_1.group(2)
        m3 = mm_1_1.group(5)
        m4 = mm_1_1.group(6)
        m5 = mm_1_1.group(7)
        m6 = mm_1_1.group(8)
        if m1 == '两' or m1 == '隔':
            m1 = '2'
        elif m1 == '每':
            m1 = '1'
        if m3 == '两' or m3 == '隔':
            m3 = '2'
        elif m3 == '每':
            m3 = '1'
        if m2 == '时':
            m2 = '小时'
        elif m2 == '日':
            m2 = '天'
        if m4 == '时':
            m4 = '小时'
        elif m4 == '日':
            m4 = '天'
        if not is_number(m1):
            sign = True
        elif not is_number(m3):
            sign = True
        if not sign:
            buf_1 = None
            buf_2 = None
            if is_number(m1) == 1:
                # 1 2 3  --->一 二 三
                m1 = int(m1)
                buf_1 = m1
                m1 = _to_chinese4(m1)
                if m1 == '二':
                    m1 = '两'
            else:
                buf_1 = unicodedata.numeric(m1)
            if is_number(m3) == 1:
                # 1 2 3  --->一 二 三
                m3 = int(m3)
                buf_2 = m3
                m3 = _to_chinese4(m3)
                if m3 == '二':
                    m3 = '两'
            else:
                buf_2 = unicodedata.numeric(m3)
            if buf_1 > buf_2:
                temp = m1
                m1 = m3
                m3 = temp
            if is_number(m5) == 2:
                # 一 二 三 ---> 1 2 3
                m5 = str(round(unicodedata.numeric(m5)))
            df_list.append(">=")
            df_list.append(m5)
            if m1 == '一':
                df_list.append(f"{m6}/{m2}")
            else:
                df_list.append(f"{m6}/{m1}{m2}")

            df_list.append(m5)
            df_list.append("且<=")
            if m3 == '一':
                df_list.append(f"{m6}/{m4}")
            else:
                df_list.append(f"{m6}/{m3}{m4}")
            buff = ''
            for temp in df_list:
                buff += str(temp)
                buff += ' '
            return buff
    if sign and mm_1_2:
        # 匹配类型：一周1次，一周1或2次，12小时给|用药一次，一日2—3次
        sign = False
        m1 = mm_1_2.group(1)
        m2 = mm_1_2.group(2)
        m3 = mm_1_2.group(3)
        m4 = mm_1_2.group(4)
        if m1 == '隔':
            m1 = '2'
        if m1 == '每':
            m1 = '1'
        if not is_number(m1):
            sign = True
        if not sign:
            if m3 == '两':
                m3 = '2'
            if m2 == '时':
                m2 = '小时'
            elif m2 == '日':
                m2 = '天'
            if is_number(m1) == 1:
                m1 = int(m1)
                m1 = _to_chinese4(m1)
                if m1 == '二':
                    m1 = '两'
            if is_number(m3) == 1:  # 1 2 3
                df_list.append("=")
                df_list.append(m3)
                if m1 == "一":
                    df_list.append(f"{m4}/{m2}")
                else:
                    df_list.append(f"{m4}/{m1}{m2}")
            elif is_number(m3) == 2:
                m3 = str(round(unicodedata.numeric(m3)))
                df_list.append("=")
                df_list.append(m3)
                if m1 == '一':
                    df_list.append(f"{m4}/{m2}")
                elif m1 == '二十四' and m2 == '小时':
                    m2 = '天'
                    df_list.append(f"{m4}/{m2}")
                else:
                    df_list.append(f"{m4}/{m1}{m2}")
            elif '-' in m3 or '～' in m3 or '~' in m3 or '一' in m3 or '或' in m3 or '—' in m3 or '至' in m3 or '－' in m3:
                m3 = re.split('～|\-|~|—|一|或|至', m3)
                df_list.append(">=")
                if is_number(m3[0]) == 2:
                    m3[0] = str(round(unicodedata.numeric(m3[0])))
                df_list.append(m3[0])
                if m1 == '一':
                    df_list.append(f"{m4}/{m2}")
                else:
                    df_list.append(f"{m4}/{m1}{m2}")
                if is_number(m3[1]) == 2:
                    m3[1] = str(round(unicodedata.numeric(m3[1])))
                df_list.append("且<=")
                df_list.append(m3[1])
                if m1 == '一':
                    df_list.append(f"{m4}/{m2}")
                else:
                    df_list.append(f"{m4}/{m1}{m2}")
            elif m3 == "数":
                df_list.append(">")
                df_list.append(1)
                df_list.append(f"{m4}/{m2}")
            else:
                sign = False
            buff = ''
            for temp in df_list:
                buff += str(temp)
                buff += ' '
            return buff
    if sign and mm_1_3:
        # 匹配类型：一周或两周1次，一或两周1次
        sign = False
        if mm_1_3.group(2) == None:
            # 匹配类型：一或两天1次
            m1 = mm_1_3.group(1)
            m2 = mm_1_3.group(4)
            m3 = mm_1_3.group(3)
            m4 = mm_1_3.group(4)
            m5 = mm_1_3.group(5)
            m6 = mm_1_3.group(6)
        else:
            # 匹配类型：一天或两天1次
            m1 = mm_1_3.group(1)
            m2 = mm_1_3.group(2)
            m3 = mm_1_3.group(3)
            m4 = mm_1_3.group(4)
            m5 = mm_1_3.group(5)
            m6 = mm_1_3.group(6)
        if m1 == '两' or m1 == '隔':
            m1 = '2'
        elif m1 == '每':
            m1 = '1'
        if m3 == '两' or m3 == '隔':
            m3 = '2'
        elif m3 == '每':
            m3 = '1'
        if m2 == '时':
            m2 = '小时'
        elif m2 == '日':
            m2 = '天'
        if m4 == '时':
            m4 = '小时'
        elif m4 == '日':
            m4 = '天'
        if not is_number(m1):
            sign = True
        elif not is_number(m3):
            sign = True
        if not sign:
            if is_number(m1) == 1:
                m1 = int(m1)
                buf_1 = m1
                m1 = _to_chinese4(m1)
                if m1 == '二':
                    m1 = '两'
            else:
                buf_1 = unicodedata.numeric(m1)
            if is_number(m3) == 1:
                m3 = int(m3)
                buf_2 = m3
                m3 = _to_chinese4(m3)
                if m3 == '二':
                    m3 = '两'
            else:
                buf_2 = unicodedata.numeric(m3)

            # 12或8小时一次，输出：>=且<= 1 次/八小时 1 次/十二小时
            if buf_1 > buf_2:
                temp = m1
                m1 = m3
                m3 = temp

            if is_number(m5) == 2:
                m5 = str(round(unicodedata.numeric(m5)))
            df_list.append(">=")
            df_list.append(m5)

            if m1 == '一':
                df_list.append(f"{m6}/{m2}")
            else:
                df_list.append(f"{m6}/{m1}{m2}")

            df_list.append("且<=")
            df_list.append(m5)
            if m3 == '一':
                df_list.append(f"{m6}/{m4}")
            else:
                df_list.append(f"{m6}/{m3}{m4}")
            buff = ''
            for temp in df_list:
                buff += str(temp)
                buff += ' '
            return buff
    if sign and mm_1_4:
        # 匹配类型：每日早晚各一次，早晚各一次
        sign = False
        m1 = mm_1_4.group(1)
        m2 = mm_1_4.group(2)
        m3 = mm_1_4.group(3)
        m4 = mm_1_4.group(4)
        if m1 == '每' or m1 == None:
            m1 = '1'
            m2 = '天'
        if not is_number(m1):
            sign = True
        if not sign:
            if m3 == '两':
                m3 = '2'
            if m2 == '日':
                m2 = '天'
            if is_number(m1) == 1:
                m1 = int(m1)
                m1 = _to_chinese4(m1)
                if m1 == '二':
                    m1 = '两'
            if is_number(m3) == 1:
                df_list.append("=")
                df_list.append(int(m3) * 2)
                if m1 == "一":
                    df_list.append(f"{m4}/{m2}")
                else:
                    df_list.append(f"{m4}/{m1}{m2}")
            elif is_number(m3) == 2:
                m3 = str(round(unicodedata.numeric(m3)))
                df_list.append("=")
                df_list.append(int(m3) * 2)
                if m1 == '一':
                    df_list.append(f"{m4}/{m2}")
                else:
                    df_list.append(f"{m4}/{m1}{m2}")
            else:
                sign = False
            buff = ''
            for temp in df_list:
                buff += str(temp)
                buff += ' '
            return buff
    if sign and mm_1_5:
        # 匹配类型：早晚各一次，早晚一次
        sign = False
        m1 = mm_1_5.group(1)
        if m1 == '两':
            m1 = 2
        if not is_number(m1):
            sign = True
        if not sign:
            if is_number(m1) == 2:
                # 一 二 三 ---> 1 2 3
                m1 = round(unicodedata.numeric(m1))
            df_list += ['=', str(m1 * 2), '次/天']
            buff = ''
            for temp in df_list:
                buff += str(temp)
                buff += ' '
            return buff
    if sign and mm_1_6:
        # 匹配类型：3次/天
        sign = False
        m1 = mm_1_6.group(1)
        if not is_number(m1):
            sign = True
        if not sign:
            if is_number(m1) == 2:
                m1 = str(round(unicodedata.numeric(m1)))
            df_list.append("=")
            df_list.append(m1)
            df_list.append("次/天")
            buff = ''
            for temp in df_list:
                buff += str(temp)
                buff += ' '
            return buff
    if sign and mm_1_7:
        # 匹配类型：每晚睡前一次，每晚一次
        sign = False
        m1 = mm_1_7.group(1)
        if not is_number(m1):
            sign = True
        if not sign:
            if is_number(m1) == 2:
                m1 = str(round(unicodedata.numeric(m1)))
            df_list.append("=")
            df_list.append(m1)
            df_list.append("次/天")
            buff = ''
            for temp in df_list:
                buff += str(temp)
                buff += ' '
            return buff
    if sign and mm_1_8:
        # 匹配类型：分早中两次服用，分早、中二次服用，分2~4次服用，分2～4次
        sign = False
        m1 = mm_1_8.group(1)
        if m1 == '两':
            m1 = '2'
        if is_number(m1) == 2:
            m1 = str(round(unicodedata.numeric(m1)))
            df_list.append("=")
            df_list.append(m1)
            df_list.append("次/天")
        elif is_number(m1) == 1:
            df_list.append("=")
            df_list.append(m1)
            df_list.append("次/天")
        elif '-' in m1 or '～' in m1 or '~' in m1 or '一' in m1 or '或' in m1 or '—' in m1 or '至' in m1 or '－' in m1:
            m1 = re.split('～|\-|~|—|一|－|或|至', m1)
            df_list.append(">=")
            if not is_number(m1[0]) or not is_number(m1[1]):
                sign = True
            if not sign:
                if is_number(m1[0]) == 2:
                    m1[0] = str(round(unicodedata.numeric(m1[0])))
                if is_number(m1[1]) == 2:
                    m1[1] = str(round(unicodedata.numeric(m1[1])))
                df_list.append(m1[0])
                df_list.append("次/天")
                df_list.append("且<=")
                df_list.append(m1[1])
                df_list.append("次/天")
        if not sign:
            buff = ''
            for temp in df_list:
                buff += str(temp)
                buff += ' '
            return buff
    if sign and mm_1_9:
        # 匹配类型：顿服，单剂，单次，分次
        sign = False
        mm = mm_1_9.group(2)
        if mm == "顿服":
            df_list.append("=")
            df_list.append("3")
            df_list.append("次/天")
        elif mm == "单剂" or mm == "单次":
            df_list.append("=")
            df_list.append("1")
            df_list.append("次/天")
        elif mm == "分次":
            df_list.append("=")
            df_list.append("2")
            df_list.append("次/天")
        if not sign:
            buff = ''
            for temp in df_list:
                buff += str(temp)
                buff += ' '
            return buff

    return None


# 药嘱.每天剂量  药嘱.每次给药剂量
def search_dayDose_eachTimeDose(mm_2_1, mm_2_2):
    sign = True
    dd_list = []
    if mm_2_1:
        # PATTERN_2_1 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+\.?\d*)(丸)?(～|\-|~|—|±|一|或者|或|至)([\u4e00-\u9fa5]|\d+\.?\d*)(丸)([\u4e00-\u9fa5]*)"
        sign = False
        str_s = mm_2_1.group(1)
        str_e = mm_2_1.group(7)
        m1 = mm_2_1.group(2)
        symbol = mm_2_1.group(4)
        m2 = mm_2_1.group(5)
        m3 = mm_2_1.group(6)
        if m1 == '俩' or m1 == "兩":
            m1 = 2
        elif m1 == '半':
            m1 = 0.5
        if m2 == '俩' or m2 == "兩":
            m2 = 2
        if not is_number(m1):
            sign = True
        elif not is_number(m2):
            sign = True
        if not sign:
            if is_number(m1) == 2:
                m1 = unicodedata.numeric(m1)
            if is_number(m2) == 2:
                m2 = unicodedata.numeric(m2)
            if float(m1) > float(m2):
                temp = m1
                m1 = m2
                m2 = temp
            if str_e.find("以上") != -1 or str_s.find("至少") != -1 or str_s.find(
                    "不少于") != -1 or str_s.find("不得少于") != -1 or str_s.find("最少") != -1 or str_s.find(
                "最低") != -1:
                dd_list.append(">=")
                dd_list.append(m1)
                dd_list.append(m3)
            elif str_s.find("不宜超过") != -1 or str_s.find("不超过") != -1 or str_s.find("不可超过") != -1 or str_s.find(
                    "不要超过") != -1 or str_s.find("不得超过") != -1 or str_s.find("最大") != -1:
                dd_list.append("<=")
                dd_list.append(m2)
                dd_list.append(m3)
            elif str_s.find("大于") != -1 or str_s.find("超过") != -1:
                dd_list.append(">")
                dd_list.append(m1)
                dd_list.append(m3)
            elif str_s.find("小于") != -1 or str_s.find("少于") != -1:
                dd_list.append("<")
                dd_list.append(m2)
                dd_list.append(m3)
            elif symbol == "±":
                dd_list.append(float(m1) - float(m2))
                dd_list.append(m3)
                dd_list.append(float(m1) + float(m2))
                dd_list.append(m3)
            else:
                dd_list.append(">=")
                dd_list.append(m1)
                dd_list.append(m3)
                dd_list.append("且<=")
                dd_list.append(m2)
                dd_list.append(m3)
            buff = ''
            for temp in dd_list:
                buff += str(temp)
                buff += ' '
            return buff, m1, m2
    if sign and mm_2_2:
        # PATTERN_2_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+\.?\d*)(丸)([\u4e00-\u9fa5]*)"
        sign = False
        str_s = mm_2_2.group(1)
        str_e = mm_2_2.group(4)
        m1 = mm_2_2.group(2)
        m2 = mm_2_2.group(3)
        if m1 == '俩' or m1 == '两':
            m1 = '二'
        elif m1 == '半':
            m1 = 0.5
        if not is_number(m1):
            sign = True
        if not sign:
            if is_number(m1) == 2:
                m1 = round(unicodedata.numeric(m1))
            if str_e.find("以上") != -1 or str_s.find("至少") != -1 or str_s.find("不少于") != -1 or str_s.find(
                    "不得少于") != -1 or str_s.find("最少") != -1 or str_s.find("最低") != -1:
                dd_list.append(">=")
                dd_list.append(m1)
                dd_list.append(m2)
            elif str_s.find("不宜超过") != -1 or str_s.find("不超过") != -1 or str_s.find("不可超过") != -1 or str_s.find(
                    "不要超过") != -1 or str_s.find("不得超过") != -1 or str_s.find("最大") != -1:
                dd_list.append("<=")
                dd_list.append(m1)
                dd_list.append(m2)
            elif str_s.find("大于") != -1 or str_s.find("超过") != -1:
                dd_list.append(">")
                dd_list.append(m1)
                dd_list.append(m2)
            elif str_s.find("小于") != -1 or str_s.find("少于") != -1:
                dd_list.append("<")
                dd_list.append(m1)
                dd_list.append(m2)
            else:
                dd_list.append("=")
                dd_list.append(m1)
                dd_list.append(m2)
            buff = ''
            for temp in dd_list:
                buff += str(temp)
                buff += ' '
            return buff, m1, None
    return None, None, None


# 药嘱.使用疗程
def search_useTreatment(mm_4_1, mm_4_2):
    ut_list = []
    sign = True
    if mm_4_1:
        # PATTERN_4_1 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(次|日|个月|月|时|天|年|周)?(～|\-|~|—|±|一|或者|或|至)([\u4e00-\u9fa5]|\d+)(次|日|个月|月|时|天|年|周)([\u4e00-\u9fa5]*)"
        sign = False
        str_s = mm_4_1.group(1)
        str_e = mm_4_1.group(7)
        m1 = mm_4_1.group(2)
        symbol = mm_4_1.group(4)
        m2 = mm_4_1.group(5)
        m3 = mm_4_1.group(6)
        if m1 == "两" or m1 == "俩":
            m1 = 2
        if m2 == "两" or m2 == "俩":
            m2 = 2
        if m3 == "日" or m3 == "次":
            m3 = '天'
        elif m3 == "个月":
            m3 = "月"
        if not is_number(m1):
            sign = True
        if not is_number(m2):
            sign = True
        if not sign:
            if is_number(m1) == 2:
                m1 = round(unicodedata.numeric(m1))
            if is_number(m2) == 2:
                m2 = round(unicodedata.numeric(m2))
            if str_e.find("以上") != -1 or str_s.find("至少") != -1 or str_s.find(
                    "不少于") != -1 or str_s.find("不得少于") != -1 or str_s.find("最少") != -1 or str_s.find(
                "最低") != -1:
                ut_list.append(">=")
                ut_list.append(m1)
                ut_list.append(m3)
            elif str_s.find("不宜超过") != -1 or str_s.find("不超过") != -1 or str_s.find("不可超过") != -1 or str_s.find(
                    "不要超过") != -1 or str_s.find("不得超过") != -1 or str_s.find("最大") != -1:
                ut_list.append("<=")
                ut_list.append(m2)
                ut_list.append(m3)
            elif str_s.find("大于") != -1 or str_s.find("超过") != -1:
                ut_list.append(">")
                ut_list.append(m1)
                ut_list.append(m3)
            elif str_s.find("小于") != -1 or str_s.find("少于") != -1:
                ut_list.append("<")
                ut_list.append(m2)
                ut_list.append(m3)
            elif symbol == "±":
                ut_list.append(float(m1) - float(m2))
                ut_list.append(m3)
                ut_list.append(float(m1) + float(m2))
                ut_list.append(m3)
            else:
                ut_list.append(">=")
                ut_list.append(m1)
                ut_list.append(m3)
                ut_list.append("且<=")
                ut_list.append(m2)
                ut_list.append(m3)
            buff = ''
            for temp in ut_list:
                buff += str(temp)
                buff += ' '
            return buff, m1, m2
    if sign and mm_4_2:
        # PATTERN_4_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(次|日|个月|月|时|天|年|周)([\u4e00-\u9fa5]*)"
        sign = False
        str_s = mm_4_2.group(1)
        str_e = mm_4_2.group(4)
        m1 = mm_4_2.group(2)
        m2 = mm_4_2.group(3)
        if m1 == "两" or m1 == "俩":
            m1 = 2
        if m2 == "日" or m2 == "次":
            m2 = '天'
        elif m2 == "个月":
            m2 = "月"
        if not is_number(m1):
            sign = True
        if not sign:
            if is_number(m1) == 2:
                m1 = round(unicodedata.numeric(m1))
            if str_e.find("以上") != -1 or str_e.find("更长") != -1 or str_s.find("至少") != -1 or str_s.find(
                    "不少于") != -1 or str_s.find(
                    "不得少于") != -1 or str_s.find("最少") != -1 or str_s.find("最低") != -1:
                ut_list.append(">=")
                ut_list.append(m1)
                ut_list.append(m2)
            elif str_s.find("不宜超过") != -1 or str_s.find("不超过") != -1 or str_s.find("不可超过") != -1 or str_s.find(
                    "不要超过") != -1 or str_s.find("不得超过") != -1 or str_s.find("最大") != -1:
                ut_list.append("<=")
                ut_list.append(m1)
                ut_list.append(m2)
            elif str_s.find("大于") != -1 or str_s.find("超过") != -1:
                ut_list.append(">")
                ut_list.append(m1)
                ut_list.append(m2)
            elif str_s.find("小于") != -1 or str_s.find("少于") != -1:
                ut_list.append("<")
                ut_list.append(m1)
                ut_list.append(m2)
            else:
                ut_list.append("=")
                ut_list.append(m1)
                ut_list.append(m2)
            buff = ''
            for temp in ut_list:
                buff += str(temp)
                buff += ' '
            return buff, m1, None
    return None, None, None


# 药嘱.每公斤每日剂量   药嘱.每公斤每次剂量
def search_doseKGEachDay_doseKGEachTime(mm_5_1, mm_5_2):
    ded_list = []
    sign = True
    if mm_5_1:
        # PATTERN_5_1 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+\.?\d*)(喷)?(～|\-|~|—|±|一|或|或者|至)([\u4e00-\u9fa5]|\d+\.?\d*)(喷)\/?k?g?([\u4e00-\u9fa5]*)"
        sign = False
        str_s = mm_5_1.group(1)
        str_e = mm_5_1.group(7)
        symbol = mm_5_1.group(4)
        m1 = mm_5_1.group(2)
        m2 = mm_5_1.group(5)
        m3 = mm_5_1.group(6)
        if m1 == '俩' or m1 == "兩":
            m1 = 2
        elif m1 == '半':
            m1 = 0.5
        if m2 == '俩' or m2 == "兩":
            m2 = 2
        if not is_number(m1):
            sign = True
        elif is_number(m1) == 2:
            m1 = unicodedata.numeric(m1)
        if not is_number(m2):
            sign = True
        elif is_number(m2) == 2:
            m2 = unicodedata.numeric(m2)
        if not sign:
            if str_e.find("以上") != -1 or str_s.find("至少") != -1 or str_s.find("不少于") != -1 or str_s.find(
                    "不得少于") != -1 or str_s.find("最少") != -1 or str_s.find("最低") != -1:
                ded_list.append(">=")
                ded_list.append(m1)
                ded_list.append(m2)
            elif str_s.find("不宜超过") != -1 or str_s.find("不超过") != -1 or str_s.find("不可超过") != -1 or str_s.find(
                    "不要超过") != -1 or str_s.find("不得超过") != -1 or str_s.find("最大") != -1 or str_s.find(
                "小于") != -1 or str_s.find("少于") != -1:
                return None  # 属于每公斤每日极量
            elif str_s.find("大于") != -1 or str_s.find("超过") != -1:
                ded_list.append(">")
                ded_list.append(m1)
                ded_list.append(m3)
            elif symbol == "±":
                ded_list.append(float(m1) - float(m2))
                ded_list.append(m3)
                ded_list.append(float(m1) + float(m2))
                ded_list.append(m3)
            else:
                ded_list.append(">=")
                ded_list.append(m1)
                ded_list.append(m3)
                ded_list.append("且<=")
                ded_list.append(m2)
                ded_list.append(m3)
            buff = ''
            for temp in ded_list:
                buff += str(temp)
                buff += ' '
            return buff
    if sign and mm_5_2:
        # PATTERN_5_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+\.?\d*)(喷)\/?k?g?([\u4e00-\u9fa5]*)"
        sign = False
        str_s = mm_5_2.group(1)
        str_e = mm_5_2.group(4)
        m1 = mm_5_2.group(2)
        m2 = mm_5_2.group(3)
        if m1 == '俩' or m1 == "兩":
            m1 = 2
        elif m1 == '半':
            m1 = 0.5
        if not is_number(m1):
            sign = True
        elif is_number(m1) == 2:
            m1 = unicodedata.numeric(m1)
            sign = True
        if not sign:
            if str_e.find("以上") != -1 or str_s.find("至少") != -1 or str_s.find("不少于") != -1 or str_s.find(
                    "不得少于") != -1 or str_s.find("最少") != -1 or str_s.find("最低") != -1:
                ded_list.append(">=")
                ded_list.append(m1)
                ded_list.append(m2)
            elif str_s.find("不宜超过") != -1 or str_s.find("不超过") != -1 or str_s.find("不可超过") != -1 or str_s.find(
                    "不要超过") != -1 or str_s.find("小于") != -1 or str_s.find(
                "少于") != -1 or str_s.find("不得超过") != -1 or str_s.find("最大") != -1:
                return None  # 属于每公斤每日极量
            elif str_s.find("小于") != -1 or str_s.find("少于") != -1:
                ded_list.append("<")
                ded_list.append(m1)
                ded_list.append(m2)
            elif str_s.find("大于") != -1 or str_s.find("超过") != -1:
                ded_list.append(">")
                ded_list.append(m1)
                ded_list.append(m2)
            else:
                ded_list.append("=")
                ded_list.append(m1)
                ded_list.append(m2)
            buff = ''
            for temp in ded_list:
                buff += str(temp)
                buff += ' '
            return buff
    return None


# 药嘱.间隔周期
def search_medicationInterva(mm_10_1, mm_10_2):
    sign = True
    drem_list = []
    if mm_10_1:
        # PATTERN_10_1 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(～|\-|~|—|±|一|或者|或|至)([\u4e00-\u9fa5]|\d+)(个月|日|年|秒|天|小时|月|分|周|时)([\u4e00-\u9fa5]*)"
        sign = False
        str_s = mm_10_1.group(1)
        str_e = mm_10_1.group(6)
        symbol = mm_10_1.group(3)
        m1 = mm_10_1.group(2)
        m2 = mm_10_1.group(4)
        m3 = mm_10_1.group(5)
        if m1 == "两" or m1 == "俩":
            m1 = 2
        if m2 == "两" or m2 == "俩":
            m2 = 2
        if not is_number(m1):
            sign = True
        elif is_number(m1) == 2:
            m1 = round(unicodedata.numeric(m1))
        if not is_number(m2):
            sign = True
        elif is_number(m2) == 2:
            m2 = round(unicodedata.numeric(m2))
        if m3 == "小时":
            m3 = "时"
        elif m3 == "日":
            m3 = "天"
        elif m3 == "个月":
            m3 = "月"
        if not sign:
            if str_e.find("以上") != -1 or str_s.find("至少") != -1 or str_s.find("不少于") != -1 or str_s.find(
                    "不得少于") != -1 or str_s.find("最少") != -1 or str_s.find("最低") != -1:
                drem_list.append(">=")
                drem_list.append(m1)
                drem_list.append(m2)
            elif str_s.find("不宜超过") != -1 or str_s.find("不超过") != -1 or str_s.find("不可超过") != -1 or str_s.find(
                    "不要超过") != -1 or str_s.find("不得超过") != -1 or str_s.find("最大") != -1:
                drem_list.append("<=")
                drem_list.append(m1)
                drem_list.append(m3)
            elif str_s.find("小于") != -1 or str_s.find("少于") != -1:
                drem_list.append("<")
                if symbol == "±":
                    drem_list.append(float(m1) + float(m2))
                else:
                    drem_list.append(m2)
                drem_list.append(m3)
            elif str_s.find("大于") != -1 or str_s.find("超过") != -1:
                drem_list.append(">")
                drem_list.append(m1)
                drem_list.append(m3)
            elif symbol == "±":
                drem_list.append(float(m1) - float(m2))
                drem_list.append(m3)
                drem_list.append(float(m1) + float(m2))
                drem_list.append(m3)
            else:
                drem_list.append(">=")
                drem_list.append(m1)
                drem_list.append(m3)
                drem_list.append("且<=")
                drem_list.append(m2)
                drem_list.append(m3)
            buff = ''
            for temp in drem_list:
                buff += str(temp)
                buff += ' '
            return buff
    if sign and mm_10_2:
        # PATTERN_10_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(个月|日|年|秒|天|小时|月|分|周|时)([\u4e00-\u9fa5]*)"
        sign = False
        str_s = mm_10_2.group(1)
        str_e = mm_10_2.group(4)
        m1 = mm_10_2.group(2)
        m2 = mm_10_2.group(3)
        if m1 == "两" or m1 == "俩":
            m1 = 2
        if not is_number(m1):
            sign = True
        elif is_number(m1) == 2:
            m1 = round(unicodedata.numeric(m1))
        if m2 == "小时":
            m2 = "时"
        elif m2 == "日":
            m2 = "天"
        elif m2 == "个月":
            m2 = "月"
        if not sign:
            if str_s.find("不得超过") != -1 or str_s.find("不超过") != -1 or str_s.find("不应超过") != -1 or str_s.find(
                    "不能超过") != -1 or str_s.find("不可超过") != -1 or str_s.find("不宜超过") != -1 or str_s.find(
                "最大") != -1 or str_s.find("不得超过") != -1:
                drem_list.append("<=")
                drem_list.append(m1)
                drem_list.append(m2)
            elif str_s.find("超过") != -1 or str_s.find("大于") != -1 or str_e.find("以上") != -1:
                drem_list.append(">")
                drem_list.append(m1)
                drem_list.append(m2)
            elif str_s.find("至少") != -1 or str_s.find("不少于") != -1 or str_s.find("不得短于") != -1 or str_s.find(
                    "不得少于") != -1 or str_s.find("不应短于") != -1 or str_s.find("不应少于") != -1 or str_s.find(
                "不可短于") != -1 or str_s.find("不可少于") != -1 or str_s.find("不宜少于") != -1 or str_s.find(
                "不宜短于") != -1:
                drem_list.append(">=")
                drem_list.append(m1)
                drem_list.append(m2)
            else:
                drem_list.append("=")
                drem_list.append(m1)
                drem_list.append(m2)
            buff = ''
            for temp in drem_list:
                buff += str(temp)
                buff += ' '
            return buff
    return None


# 药嘱.每公斤每次极量
def search_maximumDoseKGEachTime(mm_7_1, mm_7_2):
    mdet_list = []
    sign = True
    if mm_7_1:
        # PATTERN_5_1 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+\.?\d*)(喷)?(～|\-|~|—|±|一|或|或者|至)([\u4e00-\u9fa5]|\d+\.?\d*)(喷)\/?k?g?([\u4e00-\u9fa5]*)"
        sign = False
        str_s = mm_7_1.group(1)
        str_e = mm_7_1.group(7)
        symbol = mm_7_1.group(4)
        m1 = mm_7_1.group(2)
        m2 = mm_7_1.group(6)
        m3 = mm_7_1.group(5)
        if m1 == '俩' or m1 == "兩":
            m1 = 2
        elif m1 == '半':
            m1 = 0.5
        if m2 == '俩' or m2 == "兩":
            m2 = 2
        if not is_number(m1):
            sign = True
        elif not is_number(m2):
            sign = True
        if not sign:
            if is_number(m1) == 2:
                m1 = unicodedata.numeric(m1)
            if is_number(m2) == 2:
                m2 = unicodedata.numeric(m2)
            if symbol == "±":
                mdet_list.append(float(m1) + float(m2))
                mdet_list.append(m3)
            else:
                mdet_list.append("<=")
                mdet_list.append(m2)
                mdet_list.append(m3)
            buff = ''
            for temp in mdet_list:
                buff += str(temp)
                buff += ' '
            return buff
    if sign and mm_7_2:
        sign = False
        # PATTERN_5_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+\.?\d*)(喷)\/?k?g?([\u4e00-\u9fa5]*)"
        # print(mm_7_2.group(0))
        str_s = mm_7_2.group(1)
        str_e = mm_7_2.group(4)
        m1 = mm_7_2.group(2)
        m2 = mm_7_2.group(3)
        if m1 == '俩' or m1 == "兩":
            m1 = 2
        elif m1 == '半':
            m1 = 0.5
        if not is_number(m1):
            sign = True
        if not sign:
            if is_number(m1) == 2:
                m1 = unicodedata.numeric(m1)
            mdet_list.append("<=")
            mdet_list.append(m1)
            mdet_list.append(m2)
            buff = ''
            for temp in mdet_list:
                buff += str(temp)
                buff += ' '
            return buff
    return None


#  药嘱.给药目的
def search_dosingPurpose(mm_20_1):
    if mm_20_1:
        dp_list = []
        dp_list.append("=")
        dp_list.append(mm_20_1.group(2))
        buff = ''
        for temp in dp_list:
            buff += temp
            buff += ' '
        return buff
    return None


# 药嘱.持续时间
def search_duration(mm_9_1, mm_9_2):
    sign = True
    d_list = []
    # PATTERN_9_1 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(小时|分|分钟|时|秒|日|天|个月|月)?(～|\-|~|—|±|一|或者|或|至)([\u4e00-\u9fa5]|\d+)(小时|分|分钟|时|秒|日|天|个月|月)([\u4e00-\u9fa5]*)"
    if mm_9_1:
        sign = False
        str_s = mm_9_1.group(1)
        str_e = mm_9_1.group(7)
        m1 = mm_9_1.group(2)
        symbol = mm_9_1.group(4)
        m2 = mm_9_1.group(5)
        m3 = mm_9_1.group(6)
        if m3 == "小时":
            m3 = "时"
        elif m3 == "分钟":
            m3 = "分"
        elif m3 == "个月":
            m3 = "月"
        if m1 == "两" or m1 == "俩":
            m1 = 2
        if m2 == "两" or m2 == "俩":
            m2 = 2
        if not is_number(m1) or not is_number(m2):
            sign = True
        elif is_number(m1) == 2:
            m1 = round(unicodedata.numeric(m1))
        if is_number(m2) == 2:
            m2 = round(unicodedata.numeric(m2))
        if not sign:
            if str_e.find("以上") != -1 or str_s != None and str_s.find("至少") != -1 or str_s.find(
                    "不少于") != -1 or str_s.find(
                    "不得少于") != -1 or str_s.find("最少") != -1:
                d_list.append(">=")
                d_list.append(m1)
                d_list.append(m3)
            elif str_s.find("大于") != -1 or str_s.find("超过") != -1:
                d_list.append(">")
                d_list.append(m1)
                d_list.append(m3)
            elif str_s.find("不宜超过") != -1 or str_s.find("不得超过") != -1:
                d_list.append("<=")
                d_list.append(m2)
                d_list.append(m3)
            else:
                d_list.append(">=")
                d_list.append(m1)
                d_list.append(m3)
                d_list.append("且<=")
                d_list.append(m2)
                d_list.append(m3)
            buff = ''
            for temp in d_list:
                buff += str(temp)
                buff += ' '
            return buff
    # PATTERN_9_2 = r"([\u4e00-\u9fa5]*)([\u4e00-\u9fa5]|\d+)(小时|分|分钟|时|秒|日|天|个月|月)([\u4e00-\u9fa5]*)"
    if sign and mm_9_2:
        sign = False
        str_s = mm_9_2.group(1)
        str_e = mm_9_2.group(4)
        m1 = mm_9_2.group(2)
        m2 = mm_9_2.group(3)
        if m2 == "小时":
            m2 = "时"
        elif m2 == "分钟":
            m2 = "分"
        elif m2 == "个月":
            m2 = "月"
        if m1 == "两" or m1 == "俩" or m1 == "数":
            m1 = 2
        if not is_number(m1):
            sign = True
        elif is_number(m1) == 2:
            m1 = round(unicodedata.numeric(m1))
        if not sign:
            if str_e.find("以上") != -1 or str_s.find("至少") != -1 or str_s.find("不少于") != -1 or str_s.find(
                    "不得少于") != -1 or str_s.find("最少") != -1:
                d_list.append(">=")
                d_list.append(m1)
                d_list.append(m2)
            elif str_s.find("大于") != -1 or str_s.find("超过") != -1:
                d_list.append(">")
                d_list.append(m1)
                d_list.append(m2)
            elif str_s.find("不宜超过") != -1 or str_s.find("不得超过") != -1:
                d_list.append("<=")
                d_list.append(m1)
                d_list.append(m2)
            else:
                d_list.append("=")
                d_list.append(m1)
                d_list.append(m2)
            buff = ''
            for temp in d_list:
                buff += str(temp)
                buff += ' '
            return buff


#  药嘱.给药时机
def search_dosingTime(mm_30_1, mm_30_2):
    node_dosingTime = ""
    dw_list_yes = []
    dw_list_no = []
    dw_yes = []
    dw_no = []
    if mm_30_1:
        #                 print(mm_1_20.group())
        s_str = mm_30_1.group(1)
        aim_str = mm_30_1.group(2)
        if aim_str == "早饭后":
            aim_str = "早餐后"
        elif aim_str == "早上":
            aim_str = "上午"
        if s_str != None and (s_str.find('不应') == -1 or s_str.find('不许') == -1 or s_str.find('不可') == -1):
            dw_yes.append(f"{aim_str}")
        elif s_str != None and (s_str.find('不应') != -1 or s_str.find('不许') != -1 or s_str.find('不可') != -1):
            dw_no.append(f"{aim_str}")
    if mm_30_2:
        #                 print(mm_1_20.group())
        s_str = mm_30_2.group(1)
        aim_str_1 = mm_30_2.group(2)
        aim_str_2 = mm_30_2.group(3)
        if s_str != None and (s_str.find('不应') == -1 or s_str.find('不许') == -1 or s_str.find('不可') == -1):
            dw_yes.append(f"{aim_str_1}")
            dw_yes.append(f"{aim_str_2}")
        elif s_str != None and (s_str.find('不应') != -1 or s_str.find('不许') != -1 or s_str.find('不可') != -1):
            dw_no.append(f"{aim_str_1}")
            dw_no.append(f"{aim_str_2}")
    # 去重
    dw_yes = list(set(dw_yes))
    dw_no = list(set(dw_no))

    dw_str_yes = ''
    dw_str_no = ''
    first_1_20_yes = True
    first_1_20_no = True
    if len(dw_yes) != 0:
        dw_list_yes += ['属于']
        for temp in dw_yes:
            if first_1_20_yes:
                dw_str_yes += f'{temp}'
                first_1_20_yes = False
            else:
                dw_str_yes += f'|{temp}'
        dw_list_yes.append(dw_str_yes)
        buff = ''
        for temp in dw_list_yes:
            buff += str(temp)
            buff += ' '
        node_dosingTime += buff
    if len(dw_no) != 0:
        dw_list_no += ['不属于']
        for temp in dw_no:
            if first_1_20_no:
                dw_str_no += f'{temp}'
                first_1_20_no = False
            else:
                dw_str_no += f'|{temp}'
        dw_list_no.append(dw_str_no)
        buff = ''
        for temp in dw_list_no:
            buff += str(temp)
            buff += ' '
        node_dosingTime += buff
    return node_dosingTime


# 生化指标.肝/肾功能检查
def search_functionTest(mm_21_1, mm_22_1):
    nodeLiver = ""
    nodeKidney = ""
    for mm_1_21 in mm_21_1:
        kft_list = []
        if mm_1_21 and (mm_1_21.group(2) == "肝功能不全" or mm_1_21.group(2) == "肝病"):
            kft_list.append("> 80 U/L")
            buff = ''
            for temp in kft_list:
                buff += temp
            nodeLiver = buff
        elif mm_1_21 and mm_1_21.group(2) == "肝" and (mm_1_21.group(3).find("禁用") != -1 or mm_1_21.group(3).find(
                "慎用") != -1 or mm_1_21.group(1).find("严重") != -1):
            kft_list.append("> 200 U/L")
            buff = ''
            for temp in kft_list:
                buff += temp
            nodeLiver = buff
        elif mm_1_21 and mm_1_21.group(2) == "肝" and (mm_1_21.group(3).find("异常") != -1):
            kft_list.append("> 80 U/L")
            buff = ''
            for temp in kft_list:
                buff += temp
            nodeLiver = buff
    for mm_1_22 in mm_22_1:
        lft_list = []
        if mm_1_22 and (mm_1_22.group(2) == "肾功能不全" or mm_1_22.group(2) == "肾病"):
            lft_list.append("< 70 ml/min")
            buff = ''
            for temp in lft_list:
                buff += temp
            nodeKidney = buff
        elif mm_1_22 and mm_1_22.group(2) == "肾" and (mm_1_22.group(3).find("衰竭") != -1):
            lft_list.append("< 20 ml/min")
            buff = ''
            for temp in lft_list:
                buff += temp
            nodeKidney = buff
        elif mm_1_22 and mm_1_22.group(2) == "肾" and (mm_1_22.group(3).find("禁用") != -1 or mm_1_22.group(3).find(
                "慎用") != -1 or mm_1_22.group(1).find("严重") != -1):
            lft_list.append("< 30 ml/min")
            buff = ''
            for temp in lft_list:
                buff += temp
            nodeKidney = buff

        elif mm_1_22 and mm_1_22.group(2) == "肾" and (mm_1_22.group(3).find("中度异常") != -1):
            lft_list.append("< 50 ml/min")
            buff = ''
            for temp in lft_list:
                buff += temp
            nodeKidney = buff
        elif mm_1_22 and mm_1_22.group(2) == "肾" and (mm_1_22.group(3).find("异常") != -1):
            lft_list.append("< 70 ml/min")
            buff = ''
            for temp in lft_list:
                buff += temp
            nodeKidney = buff
    return nodeLiver, nodeKidney


def isDigit(tmp):
    if ord(tmp) >= 49 and ord(tmp) <= 57:
        return 1
    elif ord(tmp) >= 97 and ord(tmp) <= 122:
        return 2
    elif ord(tmp) >= 65 and ord(tmp) <= 90:
        return 3


# 字符串自增 如：220afc3eca5b11e9a1ac84ef1846dbaf ---> 220afc3eca5b11e9a1ac84ef1846dbba
def auto_increase(strOriginal):
    # strOriginal = str(strOriginal).replace("-", "")
    strReverse = strOriginal[::-1]
    strReverse = list(strReverse)
    carry = 1
    carry_sign = True
    for i in range(len(strReverse)):
        ch = strReverse[i]
        if isDigit(ch) == 1:
            if int(ch) + carry == 10:
                carry = 1
                strReverse[i] = "0"
            else:
                carry = 0
                strReverse[i] = chr(ord(ch) + 1)
                carry_sign = False
        elif isDigit(strReverse[i]) == 2:
            if ord(ch) + carry == 103:
                carry = 1
                strReverse[i] = 'a'
            else:
                carry = 0
                strReverse[i] = chr(ord(ch) + 1)
                carry_sign = False
        elif isDigit(strReverse[i]) == 3:
            if ord(ch) + carry == 10:
                carry = 1
                strReverse[i] = 'A'
            else:
                carry = 0
                strReverse[i] = chr(ord(ch) + 1)
                carry_sign = False
        if not carry_sign:
            break
    # strReverse.insert(24, '-')
    # strReverse.insert(20, '-')
    # strReverse.insert(16, '-')
    # strReverse.insert(12, '-')
    strReverse = "".join(strReverse)
    strEnd = strReverse[::-1]
    return strEnd


# 转为公司json
def jsonFormatting(original_json):
    # print(original_json)
    jsonFormat = {}
    jsonFormat["root"] = {}
    jsonFormat["template"] = "right"
    jsonFormat["theme"] = "classic"
    jsonFormat["version"] = "1.4.33"

    PATTERN_age_1 = r"(>=|>)(\d+)(岁|月|天)(且<=|且<)(\d+)(岁|月|天)"
    PATTERN_age_2 = r"(>=|>|<|<=)(\d+)(岁|月|天)"
    compiled_rule_age_1 = re.compile(PATTERN_age_1)
    compiled_rule_age_2 = re.compile(PATTERN_age_2)

    # print(json.dumps(original_json, ensure_ascii=False, sort_keys=False, indent=4, separators=(',', ': ')))
    sign = False

    # json内数据初始化
    graphid = 2
    medicineKey = "2c9180826a42e0ca016a44112a750854"  # 待取
    parentId = uuid.uuid1()
    parentId = str(parentId).replace('-', '')
    bridgeKey = uuid.uuid1()
    created = lambda: int(round(time.time() * 1000))
    for drugName in original_json:
        id = auto_increase(parentId)
        jsonFormat["root"]["data"] = {
            "text": drugName,
            "advice": None,
            "parentId": f"{parentId}",
            "drugName": drugName,
            "graph": "Start",
            "label": drugName,
            "bridgeKey": f"{bridgeKey}",
            "created": f"{created()}",
            "checkid": None,
            "medicineKey": medicineKey,
            "isEnabled": "1",
            "applyRange": None,
            "applyObject": None,
            "analysisType": None,
            "analysisResultType": None,
            "sourceId": None,
            "origMessageIds": None,
            "exprName": None,
            "expr": None,
            "express": None,
            "leftType": None,
            "rightExpr": None,
            "graphid": f"{graphid}",
            "errorType": None,
            "resultCode": None,
            "leftUnit": None,
            "leftValue": None,
            "rightUnit": None,
            "rightValue": None,
            "isDelete": "0",
            "attrid": None,
            "leftvalueid": None,
            "rightvalueid": None,
            "background": None,
            "attrName": None,
            "severity": None,
            "source": None,
            "description": "开始",
            "message": None,
            "id": f"{id}",
            "type": None,
            "expandState": "expand"
        }  # 设置第二层data 如 药品名称
        jsonFormat["root"]["children"] = []
        for key1 in original_json[f"{drugName}"]:
            for key2 in original_json[f"{drugName}"][f"{key1}"]:
                if original_json[f"{drugName}"][f"{key1}"][f"{key2}"]:
                    sign = True
                    break
            if sign == True:
                sign = False
                children = []
                jsonChildren = {}
                graphid += 1
                parentId_1 = id
                id_1 = auto_increase(parentId_1)
                jsonChildren["data"] = {
                    "text": f"{key1}",
                    "advice": None,
                    "parentId": f"{parentId_1}",
                    "drugName": None,
                    "graph": "Edge",
                    "label": f"{key1}",
                    "bridgeKey": f"{bridgeKey}",
                    "created": f"{created()}",
                    "checkid": None,
                    "medicineKey": medicineKey,
                    "isEnabled": "1",
                    "applyRange": None,
                    "applyObject": None,
                    "analysisType": None,
                    "analysisResultType": None,
                    "sourceId": None,
                    "origMessageIds": None,
                    "exprName": None,
                    "expr": None,
                    "express": None,
                    "leftType": None,
                    "rightExpr": None,
                    "graphid": f"{graphid}",
                    "errorType": None,
                    "resultCode": None,
                    "leftUnit": None,
                    "leftValue": None,
                    "rightUnit": None,
                    "rightValue": None,
                    "isDelete": "0",
                    "attrid": None,
                    "leftvalueid": None,
                    "rightvalueid": None,
                    "background": None,
                    "attrName": None,
                    "severity": None,
                    "source": None,
                    "description": None,
                    "message": None,
                    "id": f"{id_1}",
                    "type": "ROUTE",
                    "expandState": "expand"
                }
                jsonChildren["children"] = children
                for key2 in original_json[f"{drugName}"][f"{key1}"]:
                    if original_json[f"{drugName}"][f"{key1}"][f"{key2}"]:
                        if key2 != "病人资料.年龄":
                            parentId_3 = id_1
                            sign_add_children_kidney = False
                            sign_add_children_liver = False
                            if key1 == "生化指标":
                                if key2 == "生化指标.肾功能检查.内生肌酐清除率":
                                    sign_add_children_kidney = True
                                    graphid += 1
                                    parentId_2 = id_1
                                    id_2 = auto_increase(parentId_2)
                                    children_yes = []
                                    children.append(
                                        {
                                            "data": {
                                                "text": "生化指标.肾功能检查 = True",
                                                "advice": None,
                                                "parentId": f"{parentId_2}",
                                                "drugName": None,
                                                "graph": "Condition",
                                                "label": "生化指标.肾功能检查 = True",
                                                "bridgeKey": f"{bridgeKey}",
                                                "created": f"{created()}",
                                                "checkid": None,
                                                "medicineKey": medicineKey,
                                                "isEnabled": "1",
                                                "applyRange": None,
                                                "applyObject": None,
                                                "analysisType": None,
                                                "analysisResultType": None,
                                                "sourceId": None,
                                                "origMessageIds": None,
                                                "exprName": "=",
                                                "expr": "equal",
                                                "express": None,
                                                "leftType": None,
                                                "rightExpr": None,
                                                "graphid": f"{graphid}",
                                                "errorType": None,
                                                "resultCode": None,
                                                "leftUnit": None,
                                                "leftValue": "true",
                                                "rightUnit": None,
                                                "rightValue": None,
                                                "isDelete": "0",
                                                "attrid": None,
                                                "leftvalueid": None,
                                                "rightvalueid": None,
                                                "background": None,
                                                "attrName": "BiochemicalIndexes.RenalFunction",
                                                "severity": None,
                                                "source": None,
                                                "description": None,
                                                "message": None,
                                                "id": f"{id_2}",
                                                "type": None,
                                                "expandState": "expand"
                                            },
                                            "children": children_yes
                                        }
                                    )
                                    graphid += 1
                                    parentId_3 = id_2
                                    id_3 = auto_increase(parentId_3)
                                    text = "是"
                                    children_4 = []
                                    children_yes.append(
                                        {
                                            "data": {
                                                "text": f"{text}",
                                                "advice": None,
                                                "parentId": f"{parentId_3}",
                                                "drugName": None,
                                                "graph": "Edge",
                                                "label": f"{text}",
                                                "bridgeKey": f"{bridgeKey}",
                                                "created": f"{created()}",
                                                "checkid": None,
                                                "medicineKey": medicineKey,
                                                "isEnabled": "1",
                                                "applyRange": None,
                                                "applyObject": None,
                                                "analysisType": None,
                                                "analysisResultType": None,
                                                "sourceId": None,
                                                "origMessageIds": None,
                                                "exprName": None,
                                                "expr": None,
                                                "express": None,
                                                "leftType": "Boolean",
                                                "rightExpr": None,
                                                "graphid": f"{graphid}",
                                                "errorType": None,
                                                "resultCode": None,
                                                "leftUnit": None,
                                                "leftValue": "false",
                                                "rightUnit": None,
                                                "rightValue": None,
                                                "isDelete": "0",
                                                "attrid": None,
                                                "leftvalueid": None,
                                                "rightvalueid": None,
                                                "background": None,
                                                "attrName": None,
                                                "severity": None,
                                                "source": None,
                                                "description": None,
                                                "message": None,
                                                "id": f"{id_3}",
                                                "type": "EXPR"
                                            },
                                            "children": children_4
                                        }
                                    )
                                elif key2 == "生化指标.肝功能检查.谷丙转氨酶":
                                    sign_add_children_liver = True
                                    graphid += 1
                                    parentId_3 = id_2
                                    id_3 = auto_increase(parentId_3)
                                    children_yes = []
                                    children.append(
                                        {
                                            "data": {
                                                "text": "生化指标.肝功能检查 = True",
                                                "advice": None,
                                                "parentId": f"{parentId_3}",
                                                "drugName": None,
                                                "graph": "Condition",
                                                "label": "生化指标.肝功能检查 = True",
                                                "bridgeKey": f"{bridgeKey}",
                                                "created": f"{created()}",
                                                "checkid": None,
                                                "medicineKey": medicineKey,
                                                "isEnabled": "1",
                                                "applyRange": None,
                                                "applyObject": None,
                                                "analysisType": None,
                                                "analysisResultType": None,
                                                "sourceId": None,
                                                "origMessageIds": None,
                                                "exprName": "=",
                                                "expr": "equal",
                                                "express": None,
                                                "leftType": None,
                                                "rightExpr": None,
                                                "graphid": f"{graphid}",
                                                "errorType": None,
                                                "resultCode": None,
                                                "leftUnit": None,
                                                "leftValue": "true",
                                                "rightUnit": None,
                                                "rightValue": None,
                                                "isDelete": "0",
                                                "attrid": None,
                                                "leftvalueid": None,
                                                "rightvalueid": None,
                                                "background": None,
                                                "attrName": "BiochemicalIndexes.RenalFunction",
                                                "severity": None,
                                                "source": None,
                                                "description": None,
                                                "message": None,
                                                "id": f"{id_3}",
                                                "type": None,
                                                "expandState": "expand"
                                            },
                                            "children": children_yes
                                        }
                                    )
                                    graphid += 1
                                    parentId_4 = id_3
                                    id_4 = auto_increase(parentId_4)
                                    text = "是"
                                    children_4 = []
                                    children_yes.append(
                                        {
                                            "data": {
                                                "text": f"{text}",
                                                "advice": None,
                                                "parentId": f"{parentId_4}",
                                                "drugName": None,
                                                "graph": "Edge",
                                                "label": f"{text}",
                                                "bridgeKey": f"{bridgeKey}",
                                                "created": f"{created()}",
                                                "checkid": None,
                                                "medicineKey": medicineKey,
                                                "isEnabled": "1",
                                                "applyRange": None,
                                                "applyObject": None,
                                                "analysisType": None,
                                                "analysisResultType": None,
                                                "sourceId": None,
                                                "origMessageIds": None,
                                                "exprName": None,
                                                "expr": None,
                                                "express": None,
                                                "leftType": "Boolean",
                                                "rightExpr": None,
                                                "graphid": f"{graphid}",
                                                "errorType": None,
                                                "resultCode": None,
                                                "leftUnit": None,
                                                "leftValue": "false",
                                                "rightUnit": None,
                                                "rightValue": None,
                                                "isDelete": "0",
                                                "attrid": None,
                                                "leftvalueid": None,
                                                "rightvalueid": None,
                                                "background": None,
                                                "attrName": None,
                                                "severity": None,
                                                "source": None,
                                                "description": None,
                                                "message": None,
                                                "id": f"{id_4}",
                                                "type": "EXPR"
                                            },
                                            "children": children_4
                                        }
                                    )
                            id_1 = parentId_3  # 回位
                            if sign_add_children_kidney or sign_add_children_liver:
                                children_tmp = children_4
                            else:
                                children_tmp = children
                            # 加首剂使用
                            parentId_3 = id_1
                            first_use = False
                            content = original_json[f"{drugName}"][f"{key1}"][f"{key2}"]
                            if key2 == "药嘱.每次给药剂量":
                                aa = "药嘱.首剂使用 = true"
                                if content.find(aa) != -1:
                                    first_use = True
                                    content = content.replace(aa, "")
                                    graphid += 1
                                    parentId_2 = id_1
                                    id_2 = auto_increase(parentId_2)
                                    children_yes = []
                                    children.append(
                                        {
                                            "data": {
                                                "text": "药嘱.首剂使用 = true",
                                                "advice": None,
                                                "parentId": f"{parentId_2}",
                                                "drugName": None,
                                                "graph": "Condition",
                                                "label": "药嘱.首剂使用 = true",
                                                "bridgeKey": f"{bridgeKey}",
                                                "created": f"{created()}",
                                                "checkid": None,
                                                "medicineKey": medicineKey,
                                                "isEnabled": "1",
                                                "applyRange": None,
                                                "applyObject": None,
                                                "analysisType": None,
                                                "analysisResultType": None,
                                                "sourceId": None,
                                                "origMessageIds": None,
                                                "exprName": "=",
                                                "expr": "equal",
                                                "express": None,
                                                "leftType": None,
                                                "rightExpr": None,
                                                "graphid": f"{graphid}",
                                                "errorType": None,
                                                "resultCode": None,
                                                "leftUnit": None,
                                                "leftValue": "true",
                                                "rightUnit": None,
                                                "rightValue": None,
                                                "isDelete": "0",
                                                "attrid": None,
                                                "leftvalueid": None,
                                                "rightvalueid": None,
                                                "background": None,
                                                "attrName": "药嘱.首剂使用",
                                                "severity": None,
                                                "source": None,
                                                "description": None,
                                                "message": None,
                                                "id": f"{id_2}",
                                                "type": None,
                                                "expandState": "expand"
                                            },
                                            "children": children_yes
                                        }
                                    )
                                    graphid += 1
                                    parentId_3 = id_2
                                    id_4 = auto_increase(parentId_3)
                                    text = "是"
                                    children_5 = []
                                    children_yes.append(
                                        {
                                            "data": {
                                                "text": f"{text}",
                                                "advice": None,
                                                "parentId": f"{parentId_3}",
                                                "drugName": None,
                                                "graph": "Edge",
                                                "label": f"{text}",
                                                "bridgeKey": f"{bridgeKey}",
                                                "created": f"{created()}",
                                                "checkid": None,
                                                "medicineKey": medicineKey,
                                                "isEnabled": "1",
                                                "applyRange": None,
                                                "applyObject": None,
                                                "analysisType": None,
                                                "analysisResultType": None,
                                                "sourceId": None,
                                                "origMessageIds": None,
                                                "exprName": None,
                                                "expr": None,
                                                "express": None,
                                                "leftType": "Boolean",
                                                "rightExpr": None,
                                                "graphid": f"{graphid}",
                                                "errorType": None,
                                                "resultCode": None,
                                                "leftUnit": None,
                                                "leftValue": "false",
                                                "rightUnit": None,
                                                "rightValue": None,
                                                "isDelete": "0",
                                                "attrid": None,
                                                "leftvalueid": None,
                                                "rightvalueid": None,
                                                "background": None,
                                                "attrName": None,
                                                "severity": None,
                                                "source": None,
                                                "description": None,
                                                "message": None,
                                                "id": f"{id_3}",
                                                "type": "EXPR"
                                            },
                                            "children": children_5
                                        }
                                    )
                            if first_use:
                                children_tmp = children_5
                            id_1 = parentId_3  # 回位
                            # print(key2)
                            # print(content)
                            if key2 == "药嘱.给药频率":  # 当把所有字段都存储为[]时(不止是药嘱.给药频率)，不需要有else部分
                                for content in original_json[f"{drugName}"][f"{key1}"][f"{key2}"]:
                                    text = key2 + " " + content
                                    # 分割填空
                                    str_list = content.split()
                                    exprName = None
                                    leftValue = None
                                    leftUnit = None
                                    rightValue = None
                                    rightUnit = None
                                    if len(str_list) > 3:
                                        exprName = str_list[0] + str_list[3]
                                        leftValue = str_list[1]
                                        leftUnit = str_list[2]
                                        rightValue = str_list[4]
                                        rightUnit = str_list[5]
                                    elif len(str_list) == 3:
                                        exprName = str_list[0]
                                        leftValue = str_list[1]
                                        leftUnit = str_list[2]
                                        rightValue = None
                                        rightUnit = None
                                    elif len(str_list) == 2:
                                        exprName = str_list[0]
                                        leftValue = str_list[1]
                                        leftUnit = None
                                        rightValue = None
                                        rightUnit = None
                                    graphid += 1
                                    parentId_2 = id_1
                                    id_2 = auto_increase(parentId_2)
                                    expr = None
                                    if exprName == "属于":
                                        expr = "belong"
                                    elif exprName == "=":
                                        expr = "equal"
                                    elif exprName == ">=且<=":
                                        expr = "greaterThanEqualAndLessEqual"
                                    elif exprName == ">=":
                                        expr = "greaterThanEqual"
                                    elif exprName == "<=":
                                        expr = "lessThanEqual"
                                    elif exprName == ">":
                                        expr = "greaterThan"
                                    elif exprName == "<":
                                        expr = "lessThan"
                                    children_no = []
                                    children_tmp.append(
                                        {
                                            "data": {
                                                "text": text,
                                                "advice": None,
                                                "parentId": f"{parentId_2}",
                                                "drugName": None,
                                                "graph": "condition",
                                                "label": text,
                                                "bridgeKey": f"{bridgeKey}",
                                                "created": f"{created()}",
                                                "checkid": None,
                                                "medicineKey": medicineKey,
                                                "isEnabled": "1",
                                                "applyRange": None,
                                                "applyObject": None,
                                                "analysisType": None,
                                                "analysisResultType": None,
                                                "sourceId": None,
                                                "origMessageIds": None,
                                                "exprName": exprName,
                                                "expr": expr,
                                                "express": None,
                                                "leftType": None,
                                                "rightExpr": None,
                                                "graphid": f"{graphid}",
                                                "errorType": None,
                                                "resultCode": None,
                                                "leftUnit": leftUnit,
                                                "leftValue": leftValue,
                                                "rightUnit": rightUnit,
                                                "rightValue": rightValue,
                                                "isDelete": "0",
                                                "attrid": None,
                                                "leftvalueid": None,
                                                "rightvalueid": None,
                                                "background": None,
                                                "attrName": key2,  # 需要转化为英文？
                                                "severity": None,
                                                "source": None,
                                                "description": None,
                                                "message": None,
                                                "id": f"{id_2}",
                                                "type": "ROUTE",
                                                "expandState": "expand"
                                            },
                                            "children": children_no
                                        }
                                    )
                            else:
                                text = key2 + " " + content
                                # 分割填空
                                str_list = content.split()
                                exprName = None
                                leftValue = None
                                leftUnit = None
                                rightValue = None
                                rightUnit = None
                                if len(str_list) > 3:
                                    exprName = str_list[0] + str_list[3]
                                    leftValue = str_list[1]
                                    leftUnit = str_list[2]
                                    rightValue = str_list[4]
                                    rightUnit = str_list[5]
                                elif len(str_list) == 3:
                                    exprName = str_list[0]
                                    leftValue = str_list[1]
                                    leftUnit = str_list[2]
                                    rightValue = None
                                    rightUnit = None
                                elif len(str_list) == 2:
                                    exprName = str_list[0]
                                    leftValue = str_list[1]
                                    leftUnit = None
                                    rightValue = None
                                    rightUnit = None
                                graphid += 1
                                parentId_2 = id_1
                                id_2 = auto_increase(parentId_2)
                                expr = None
                                if exprName == "属于":
                                    expr = "belong"
                                elif exprName == "=":
                                    expr = "equal"
                                elif exprName == ">=且<=":
                                    expr = "greaterThanEqualAndLessEqual"
                                elif exprName == ">=":
                                    expr = "greaterThanEqual"
                                elif exprName == "<=":
                                    expr = "lessThanEqual"
                                elif exprName == ">":
                                    expr = "greaterThan"
                                elif exprName == "<":
                                    expr = "lessThan"
                                children_no = []
                                children_tmp.append(
                                    {
                                        "data": {
                                            "text": text,
                                            "advice": None,
                                            "parentId": f"{parentId_2}",
                                            "drugName": None,
                                            "graph": "condition",
                                            "label": text,
                                            "bridgeKey": f"{bridgeKey}",
                                            "created": f"{created()}",
                                            "checkid": None,
                                            "medicineKey": medicineKey,
                                            "isEnabled": "1",
                                            "applyRange": None,
                                            "applyObject": None,
                                            "analysisType": None,
                                            "analysisResultType": None,
                                            "sourceId": None,
                                            "origMessageIds": None,
                                            "exprName": exprName,
                                            "expr": expr,
                                            "express": None,
                                            "leftType": None,
                                            "rightExpr": None,
                                            "graphid": f"{graphid}",
                                            "errorType": None,
                                            "resultCode": None,
                                            "leftUnit": leftUnit,
                                            "leftValue": leftValue,
                                            "rightUnit": rightUnit,
                                            "rightValue": rightValue,
                                            "isDelete": "0",
                                            "attrid": None,
                                            "leftvalueid": None,
                                            "rightvalueid": None,
                                            "background": None,
                                            "attrName": key2,  # 需要转化为英文？
                                            "severity": None,
                                            "source": None,
                                            "description": None,
                                            "message": None,
                                            "id": f"{id_2}",
                                            "type": "ROUTE",
                                            "expandState": "expand"
                                        },
                                        "children": children_no
                                    }
                                )
                            children_advice = []
                            if key1 == "用法用量" or key1 == "给药途径":
                                text = "否"
                                graphid += 1
                                parentId_3 = id_2
                                id_3 = auto_increase(parentId_3)
                                children_advice.append(
                                    {
                                        "data": {
                                            "text": "一、消息依据：说明书\n二、警示等级：3级\n三、错误描述：\n\n四、药品建议\n\n",
                                            "advice": None,
                                            "parentId": f"{parentId_3}",
                                            "drugName": None,
                                            "graph": "Result",
                                            "label": None,
                                            "bridgeKey": f"{bridgeKey}",
                                            "created": f"{created()}",
                                            "checkid": None,
                                            "medicineKey": medicineKey,
                                            "isEnabled": "1",
                                            "applyRange": None,
                                            "applyObject": None,
                                            "analysisType": None,
                                            "analysisResultType": None,
                                            "sourceId": None,
                                            "origMessageIds": None,
                                            "exprName": None,
                                            "expr": None,
                                            "express": None,
                                            "leftType": None,
                                            "rightExpr": None,
                                            "graphid": f"{graphid}",
                                            "errorType": f"{key1}",
                                            "resultCode": None,
                                            "leftUnit": None,
                                            "leftValue": None,
                                            "rightUnit": None,
                                            "rightValue": None,
                                            "isDelete": "0",
                                            "attrid": None,
                                            "leftvalueid": None,
                                            "rightvalueid": None,
                                            "background": "#92d050",
                                            "attrName": None,
                                            "severity": "3",
                                            "source": "说明书",
                                            "description": None,
                                            "message": None,
                                            "id": f"{id_3}",
                                            "type": "ROUTE",
                                            "expandState": "expand"
                                        },
                                        "children": []
                                    }
                                )
                                id_2 = id_3
                            else:
                                text = "是"
                                graphid += 1
                                parentId_3 = id_2
                                id_3 = auto_increase(parentId_3)
                                children_advice.append(
                                    {
                                        "data": {
                                            "text": "一、消息依据：说明书\n二、警示等级：3级\n三、错误描述：\n\n四、药品建议\n\n",
                                            "advice": None,
                                            "parentId": f"{parentId_3}",
                                            "drugName": None,
                                            "graph": "Result",
                                            "label": None,
                                            "bridgeKey": f"{bridgeKey}",
                                            "created": f"{created()}",
                                            "checkid": None,
                                            "medicineKey": medicineKey,
                                            "isEnabled": "1",
                                            "applyRange": None,
                                            "applyObject": None,
                                            "analysisType": None,
                                            "analysisResultType": None,
                                            "sourceId": None,
                                            "origMessageIds": None,
                                            "exprName": None,
                                            "expr": None,
                                            "express": None,
                                            "leftType": None,
                                            "rightExpr": None,
                                            "graphid": f"{graphid}",
                                            "errorType": f"{key1}",
                                            "resultCode": None,
                                            "leftUnit": None,
                                            "leftValue": None,
                                            "rightUnit": None,
                                            "rightValue": None,
                                            "isDelete": "0",
                                            "attrid": None,
                                            "leftvalueid": None,
                                            "rightvalueid": None,
                                            "background": "#FF4E4E",
                                            "attrName": None,
                                            "severity": "7",
                                            "source": "说明书",
                                            "description": None,
                                            "message": None,
                                            "id": f"{id_3}",
                                            "type": "ROUTE",
                                            "expandState": "expand"
                                        },
                                        "children": []
                                    }
                                )
                                id_2 = id_3
                            graphid += 1
                            parentId_3 = id_2
                            id_3 = auto_increase(parentId_3)
                            children_no.append(
                                {
                                    "data": {
                                        "text": f"{text}",
                                        "advice": None,
                                        "parentId": f"{parentId_3}",
                                        "drugName": None,
                                        "graph": "Edge",
                                        "label": f"{text}",
                                        "bridgeKey": f"{bridgeKey}",
                                        "created": f"{created()}",
                                        "checkid": None,
                                        "medicineKey": medicineKey,
                                        "isEnabled": "1",
                                        "applyRange": None,
                                        "applyObject": None,
                                        "analysisType": None,
                                        "analysisResultType": None,
                                        "sourceId": None,
                                        "origMessageIds": None,
                                        "exprName": None,
                                        "expr": None,
                                        "express": None,
                                        "leftType": "Boolean",
                                        "rightExpr": None,
                                        "graphid": f"{graphid}",
                                        "errorType": None,
                                        "resultCode": None,
                                        "leftUnit": None,
                                        "leftValue": "false",
                                        "rightUnit": None,
                                        "rightValue": None,
                                        "isDelete": "0",
                                        "attrid": None,
                                        "leftvalueid": None,
                                        "rightvalueid": None,
                                        "background": None,
                                        "attrName": None,
                                        "severity": None,
                                        "source": None,
                                        "description": None,
                                        "message": None,
                                        "id": f"{id_3}",
                                        "type": "EXPR"
                                    },
                                    "children": children_advice
                                }
                            )
                        elif key2 == "病人资料.年龄":
                            if not original_json[f"{drugName}"][f"{key1}"][f"{key2}"]:
                                continue
                            children2 = []
                            graphid += 1
                            parentId_2 = id_1
                            id_2 = auto_increase(parentId_2)
                            children.append(
                                {
                                    "data": {
                                        "text": key2,
                                        "advice": None,
                                        "parentId": f"{parentId_2}",
                                        "drugName": None,
                                        "graph": "Edge",
                                        "label": key2,
                                        "bridgeKey": f"{bridgeKey}",
                                        "created": f"{created()}",
                                        "checkid": None,
                                        "medicineKey": medicineKey,
                                        "isEnabled": "1",
                                        "applyRange": None,
                                        "applyObject": None,
                                        "analysisType": None,
                                        "analysisResultType": None,
                                        "sourceId": None,
                                        "origMessageIds": None,
                                        "exprName": None,
                                        "expr": None,
                                        "express": None,
                                        "leftType": None,
                                        "rightExpr": None,
                                        "graphid": f"{graphid}",
                                        "errorType": None,
                                        "resultCode": None,
                                        "leftUnit": None,
                                        "leftValue": None,
                                        "rightUnit": None,
                                        "rightValue": None,
                                        "isDelete": "0",
                                        "attrid": None,
                                        "leftvalueid": None,
                                        "rightvalueid": None,
                                        "background": None,
                                        "attrName": None,
                                        "severity": None,
                                        "source": None,
                                        "description": None,
                                        "message": None,
                                        "id": f"{id_2}",
                                        "type": "ROUTE",
                                        "expandState": "expand"
                                    },
                                    "children": children2
                                }
                            )
                            for key3 in original_json[f"{drugName}"][f"{key1}"][f"{key2}"]:
                                ## 若else没有数据，则不显示
                                # if key3 == "else":
                                #     key3_exist_sign = False
                                #     for key4 in original_json[f"{drugName}"][f"{key1}"][f"{key2}"][f"{key3}"]:
                                #         if original_json[f"{drugName}"][f"{key1}"][f"{key2}"][f"{key3}"][f"{key4}"]:
                                #             key3_exist_sign = True  # 表示应该有else
                                #             break
                                #     if not key3_exist_sign:
                                #         continue
                                children3 = []
                                # 分割填空
                                str_age = key3.replace(' ', '')
                                mm_age_1 = re.search(compiled_rule_age_1, str_age)
                                mm_age_2 = re.search(compiled_rule_age_2, str_age)
                                exprName = None
                                leftValue = None
                                leftUnit = None
                                rightValue = None
                                rightUnit = None
                                if mm_age_1:
                                    # >= 2 月 且< 2 岁
                                    exprName = mm_age_1.group(1) + mm_age_1.group(4)
                                    leftValue = mm_age_1.group(2)
                                    leftUnit = mm_age_1.group(3)
                                    rightValue = mm_age_1.group(5)
                                    rightUnit = mm_age_1.group(6)
                                elif mm_age_2:
                                    # < 28 天
                                    exprName = mm_age_2.group(1)
                                    leftValue = mm_age_2.group(2)
                                    leftUnit = mm_age_2.group(3)
                                    rightValue = None
                                    rightUnit = None
                                expr = None
                                if exprName == "属于":
                                    expr = "belong"
                                elif exprName == "=":
                                    expr = "equal"
                                elif exprName == ">=且<=":
                                    expr = "greaterThanEqualAndLessEqual"
                                elif exprName == ">=":
                                    expr = "greaterThanEqual"
                                elif exprName == "<=":
                                    expr = "lessThanEqual"
                                elif exprName == ">":
                                    expr = "greaterThan"
                                elif exprName == "<":
                                    expr = "lessThan"
                                graphid += 1
                                parentId_3 = id_2
                                id_3 = auto_increase(parentId_3)
                                children2.append(
                                    {
                                        "data": {
                                            "text": key3,
                                            "advice": None,
                                            "parentId": f"{parentId_3}",
                                            "drugName": None,
                                            "graph": "Edge",
                                            "label": key3,
                                            "bridgeKey": f"{bridgeKey}",
                                            "created": f"{created()}",
                                            "checkid": None,
                                            "medicineKey": medicineKey,
                                            "isEnabled": "1",
                                            "applyRange": None,
                                            "applyObject": None,
                                            "analysisType": None,
                                            "analysisResultType": None,
                                            "sourceId": None,
                                            "origMessageIds": None,
                                            "exprName": exprName,
                                            "expr": expr,
                                            "express": None,
                                            "leftType": None,
                                            "rightExpr": None,
                                            "graphid": f"{graphid}",
                                            "errorType": None,
                                            "resultCode": None,
                                            "leftUnit": leftUnit,
                                            "leftValue": leftValue,
                                            "rightUnit": rightUnit,
                                            "rightValue": rightValue,
                                            "isDelete": "0",
                                            "attrid": None,
                                            "leftvalueid": None,
                                            "rightvalueid": None,
                                            "background": None,
                                            "attrName": None,
                                            "severity": None,
                                            "source": None,
                                            "description": None,
                                            "message": None,
                                            "id": f"{id_3}",
                                            "type": "ROUTE",
                                            "expandState": "expand"
                                        },
                                        "children": children3
                                    }
                                )
                                for key4 in original_json[f"{drugName}"][f"{key1}"][f"{key2}"][f"{key3}"]:
                                    if original_json[f"{drugName}"][f"{key1}"][f"{key2}"][f"{key3}"][f"{key4}"]:
                                        if key4 == "药嘱.给药频率":  # 当把所有字段都存储为[]时(不止是药嘱.给药频率)，不需要有else部分
                                            for tmp in original_json[f"{drugName}"][f"{key1}"][f"{key2}"][f"{key3}"][
                                                f"{key4}"]:
                                                text = key4 + " " + tmp

                                                # 分割填空
                                                str1 = tmp
                                                str_list = str1.split()
                                                # print(str_list)
                                                exprName = None
                                                leftValue = None
                                                leftUnit = None
                                                rightValue = None
                                                rightUnit = None
                                                expr = None
                                                if len(str_list) > 3:
                                                    exprName = str_list[0] + str_list[3]
                                                    leftValue = str_list[1]
                                                    leftUnit = str_list[2]
                                                    rightValue = str_list[4]
                                                    rightUnit = str_list[5]
                                                elif len(str_list) == 3:
                                                    exprName = str_list[0]
                                                    leftValue = str_list[1]
                                                    leftUnit = str_list[2]
                                                    rightValue = None
                                                    rightUnit = None
                                                elif len(str_list) == 2:
                                                    exprName = str_list[0]
                                                    leftValue = str_list[1]
                                                    leftUnit = None
                                                    rightValue = None
                                                    rightUnit = None
                                                graphid += 1
                                                parentId_4 = id_3
                                                id_4 = auto_increase(parentId_4)
                                                if exprName == "属于":
                                                    expr = "belong"
                                                elif exprName == "=":
                                                    expr = "equal"
                                                elif exprName == ">=且<=":
                                                    expr = "greaterThanEqualAndLessEqual"
                                                elif exprName == ">=":
                                                    expr = "greaterThanEqual"
                                                elif exprName == "<=":
                                                    expr = "lessThanEqual"
                                                elif exprName == ">":
                                                    expr = "greaterThan"
                                                elif exprName == "<":
                                                    expr = "lessThan"
                                                children_no = []
                                                children3.append(
                                                    {
                                                        "data": {
                                                            "text": text,
                                                            "advice": None,
                                                            "parentId": f"{parentId_4}",
                                                            "drugName": None,
                                                            "graph": "Condition",
                                                            "label": text,
                                                            "bridgeKey": f"{bridgeKey}",
                                                            "created": f"{created()}",
                                                            "checkid": None,
                                                            "medicineKey": medicineKey,
                                                            "isEnabled": "1",
                                                            "applyRange": None,
                                                            "applyObject": None,
                                                            "analysisType": None,
                                                            "analysisResultType": None,
                                                            "sourceId": None,
                                                            "origMessageIds": None,
                                                            "exprName": exprName,
                                                            "expr": expr,
                                                            "express": None,
                                                            "leftType": None,
                                                            "rightExpr": None,
                                                            "graphid": f"{graphid}",
                                                            "errorType": None,
                                                            "resultCode": None,
                                                            "leftUnit": leftUnit,
                                                            "leftValue": leftValue,
                                                            "rightUnit": rightUnit,
                                                            "rightValue": rightValue,
                                                            "isDelete": "0",
                                                            "attrid": None,
                                                            "leftvalueid": None,
                                                            "rightvalueid": None,
                                                            "background": None,
                                                            "attrName": key4,  # 需要转化为英文？
                                                            "severity": None,
                                                            "source": None,
                                                            "description": None,
                                                            "message": text,
                                                            "id": f"{id_4}",
                                                            "type": "ROUTE",
                                                            "expandState": "expand"
                                                        },
                                                        "children": children_no
                                                    }
                                                )
                                                children_advice = []
                                                if key1 == "用法用量" or key1 == "给药途径":  # 警示级颜色？
                                                    text = "否"
                                                    graphid += 1
                                                    parentId_5 = id_4
                                                    id_5 = auto_increase(parentId_5)
                                                    children_advice.append(
                                                        {
                                                            "data": {
                                                                "text": "一、消息依据：说明书\n二、警示等级：3级\n三、错误描述：\n\n四、药品建议\n\n",
                                                                "advice": None,
                                                                "parentId": f"{parentId_5}",
                                                                "drugName": None,
                                                                "graph": "Result",
                                                                "label": None,
                                                                "bridgeKey": f"{bridgeKey}",
                                                                "created": f"{created()}",
                                                                "checkid": None,
                                                                "medicineKey": medicineKey,
                                                                "isEnabled": "1",
                                                                "applyRange": None,
                                                                "applyObject": None,
                                                                "analysisType": None,
                                                                "analysisResultType": None,
                                                                "sourceId": None,
                                                                "origMessageIds": None,
                                                                "exprName": None,
                                                                "expr": None,
                                                                "express": None,
                                                                "leftType": None,
                                                                "rightExpr": None,
                                                                "graphid": f"{graphid}",
                                                                "errorType": f"{key1}",
                                                                "resultCode": None,
                                                                "leftUnit": None,
                                                                "leftValue": None,
                                                                "rightUnit": None,
                                                                "rightValue": None,
                                                                "isDelete": "0",
                                                                "attrid": None,
                                                                "leftvalueid": None,
                                                                "rightvalueid": None,
                                                                "background": "#92d050",
                                                                "attrName": None,
                                                                "severity": "3",
                                                                "source": "说明书",
                                                                "description": None,
                                                                "message": None,
                                                                "id": f"{id_5}",
                                                                "type": "ROUTE",
                                                                "expandState": "expand"
                                                            },
                                                            "children": []
                                                        }
                                                    )
                                                    id_4 = id_5
                                                else:
                                                    text = "是"
                                                graphid += 1
                                                parentId_5 = id_4
                                                id_5 = auto_increase(parentId_5)
                                                children_no.append(
                                                    {
                                                        "data": {
                                                            "text": f"{text}",
                                                            "advice": None,
                                                            "parentId": f"{parentId_5}",
                                                            "drugName": None,
                                                            "graph": "Edge",
                                                            "label": f"{text}",
                                                            "bridgeKey": f"{bridgeKey}",
                                                            "created": f"{created()}",
                                                            "checkid": None,
                                                            "medicineKey": medicineKey,
                                                            "isEnabled": "1",
                                                            "applyRange": None,
                                                            "applyObject": None,
                                                            "analysisType": None,
                                                            "analysisResultType": None,
                                                            "sourceId": None,
                                                            "origMessageIds": None,
                                                            "exprName": None,
                                                            "expr": None,
                                                            "express": None,
                                                            "leftType": "Boolean",
                                                            "rightExpr": None,
                                                            "graphid": f"{graphid}",
                                                            "errorType": None,
                                                            "resultCode": None,
                                                            "leftUnit": None,
                                                            "leftValue": "false",
                                                            "rightUnit": None,
                                                            "rightValue": None,
                                                            "isDelete": "0",
                                                            "attrid": None,
                                                            "leftvalueid": None,
                                                            "rightvalueid": None,
                                                            "background": None,
                                                            "attrName": None,
                                                            "severity": None,
                                                            "source": None,
                                                            "description": None,
                                                            "message": None,
                                                            "id": f"{id_5}",
                                                            "type": "EXPR"
                                                        },
                                                        "children": children_advice
                                                    }
                                                )
                                        else:
                                            text = key4 + " " + \
                                                   original_json[f"{drugName}"][f"{key1}"][f"{key2}"][f"{key3}"][
                                                       f"{key4}"]

                                            # 分割填空
                                            str1 = original_json[f"{drugName}"][f"{key1}"][f"{key2}"][f"{key3}"][
                                                f"{key4}"]
                                            str_list = str1.split()
                                            # print(str_list)
                                            exprName = None
                                            leftValue = None
                                            leftUnit = None
                                            rightValue = None
                                            rightUnit = None
                                            expr = None
                                            if len(str_list) > 3:
                                                exprName = str_list[0] + str_list[3]
                                                leftValue = str_list[1]
                                                leftUnit = str_list[2]
                                                rightValue = str_list[4]
                                                rightUnit = str_list[5]
                                            elif len(str_list) == 3:
                                                exprName = str_list[0]
                                                leftValue = str_list[1]
                                                leftUnit = str_list[2]
                                                rightValue = None
                                                rightUnit = None
                                            elif len(str_list) == 2:
                                                exprName = str_list[0]
                                                leftValue = str_list[1]
                                                leftUnit = None
                                                rightValue = None
                                                rightUnit = None
                                            graphid += 1
                                            parentId_4 = id_3
                                            id_4 = auto_increase(parentId_4)
                                            if exprName == "属于":
                                                expr = "belong"
                                            elif exprName == "=":
                                                expr = "equal"
                                            elif exprName == ">=且<=":
                                                expr = "greaterThanEqualAndLessEqual"
                                            elif exprName == ">=":
                                                expr = "greaterThanEqual"
                                            elif exprName == "<=":
                                                expr = "lessThanEqual"
                                            elif exprName == ">":
                                                expr = "greaterThan"
                                            elif exprName == "<":
                                                expr = "lessThan"
                                            children_no = []
                                            children3.append(
                                                {
                                                    "data": {
                                                        "text": text,
                                                        "advice": None,
                                                        "parentId": f"{parentId_4}",
                                                        "drugName": None,
                                                        "graph": "Condition",
                                                        "label": text,
                                                        "bridgeKey": f"{bridgeKey}",
                                                        "created": f"{created()}",
                                                        "checkid": None,
                                                        "medicineKey": medicineKey,
                                                        "isEnabled": "1",
                                                        "applyRange": None,
                                                        "applyObject": None,
                                                        "analysisType": None,
                                                        "analysisResultType": None,
                                                        "sourceId": None,
                                                        "origMessageIds": None,
                                                        "exprName": exprName,
                                                        "expr": expr,
                                                        "express": None,
                                                        "leftType": None,
                                                        "rightExpr": None,
                                                        "graphid": f"{graphid}",
                                                        "errorType": None,
                                                        "resultCode": None,
                                                        "leftUnit": leftUnit,
                                                        "leftValue": leftValue,
                                                        "rightUnit": rightUnit,
                                                        "rightValue": rightValue,
                                                        "isDelete": "0",
                                                        "attrid": None,
                                                        "leftvalueid": None,
                                                        "rightvalueid": None,
                                                        "background": None,
                                                        "attrName": key4,  # 需要转化为英文？
                                                        "severity": None,
                                                        "source": None,
                                                        "description": None,
                                                        "message": text,
                                                        "id": f"{id_4}",
                                                        "type": "ROUTE",
                                                        "expandState": "expand"
                                                    },
                                                    "children": children_no
                                                }
                                            )
                                            children_advice = []
                                            if key1 == "用法用量" or key1 == "给药途径":  # 警示级颜色？
                                                text = "否"
                                                graphid += 1
                                                parentId_5 = id_4
                                                id_5 = auto_increase(parentId_5)
                                                children_advice.append(
                                                    {
                                                        "data": {
                                                            "text": "一、消息依据：说明书\n二、警示等级：3级\n三、错误描述：\n\n四、药品建议\n\n",
                                                            "advice": None,
                                                            "parentId": f"{parentId_5}",
                                                            "drugName": None,
                                                            "graph": "Result",
                                                            "label": None,
                                                            "bridgeKey": f"{bridgeKey}",
                                                            "created": f"{created()}",
                                                            "checkid": None,
                                                            "medicineKey": medicineKey,
                                                            "isEnabled": "1",
                                                            "applyRange": None,
                                                            "applyObject": None,
                                                            "analysisType": None,
                                                            "analysisResultType": None,
                                                            "sourceId": None,
                                                            "origMessageIds": None,
                                                            "exprName": None,
                                                            "expr": None,
                                                            "express": None,
                                                            "leftType": None,
                                                            "rightExpr": None,
                                                            "graphid": f"{graphid}",
                                                            "errorType": f"{key1}",
                                                            "resultCode": None,
                                                            "leftUnit": None,
                                                            "leftValue": None,
                                                            "rightUnit": None,
                                                            "rightValue": None,
                                                            "isDelete": "0",
                                                            "attrid": None,
                                                            "leftvalueid": None,
                                                            "rightvalueid": None,
                                                            "background": "#92d050",
                                                            "attrName": None,
                                                            "severity": "3",
                                                            "source": "说明书",
                                                            "description": None,
                                                            "message": None,
                                                            "id": f"{id_5}",
                                                            "type": "ROUTE",
                                                            "expandState": "expand"
                                                        },
                                                        "children": []
                                                    }
                                                )
                                                id_4 = id_5
                                            else:
                                                text = "是"
                                            graphid += 1
                                            parentId_5 = id_4
                                            id_5 = auto_increase(parentId_5)
                                            children_no.append(
                                                {
                                                    "data": {
                                                        "text": f"{text}",
                                                        "advice": None,
                                                        "parentId": f"{parentId_5}",
                                                        "drugName": None,
                                                        "graph": "Edge",
                                                        "label": f"{text}",
                                                        "bridgeKey": f"{bridgeKey}",
                                                        "created": f"{created()}",
                                                        "checkid": None,
                                                        "medicineKey": medicineKey,
                                                        "isEnabled": "1",
                                                        "applyRange": None,
                                                        "applyObject": None,
                                                        "analysisType": None,
                                                        "analysisResultType": None,
                                                        "sourceId": None,
                                                        "origMessageIds": None,
                                                        "exprName": None,
                                                        "expr": None,
                                                        "express": None,
                                                        "leftType": "Boolean",
                                                        "rightExpr": None,
                                                        "graphid": f"{graphid}",
                                                        "errorType": None,
                                                        "resultCode": None,
                                                        "leftUnit": None,
                                                        "leftValue": "false",
                                                        "rightUnit": None,
                                                        "rightValue": None,
                                                        "isDelete": "0",
                                                        "attrid": None,
                                                        "leftvalueid": None,
                                                        "rightvalueid": None,
                                                        "background": None,
                                                        "attrName": None,
                                                        "severity": None,
                                                        "source": None,
                                                        "description": None,
                                                        "message": None,
                                                        "id": f"{id_5}",
                                                        "type": "EXPR"
                                                    },
                                                    "children": children_advice
                                                }
                                            )
                # 设置第三层 data如用法用量 和 child
                jsonFormat["root"]["children"].append(jsonChildren)  # 设置第二层child
        # print(json.dumps(jsonFormat, ensure_ascii=False, sort_keys=False, indent=4, separators=(',', ': ')))
    return jsonFormat


if __name__ == '__main__':
    docfile = 'M10042356.doc'
    read_docx(docfile)
    filed_dict, drug_json = extract_one_sample(docfile)
    print(filed_dict)
    print(drug_json)
